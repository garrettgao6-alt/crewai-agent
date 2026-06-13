import json
from io import BytesIO
import os
from xml.sax.saxutils import escape
import time

import requests
import streamlit as st

import project_store
import prompt_store
import subscription_store
import template_store
import user_store


API_URL = os.getenv("API_URL", "http://127.0.0.1:8000/analyze")
REQUEST_TIMEOUT_SECONDS = 60
VERSION = "1.0"
HISTORY_LIMIT = 10
DOCUMENT_TEXT_LIMIT = 12000
DOCUMENT_TOTAL_TEXT_LIMIT = 30000
DOCUMENT_ANALYSIS_TYPES = [
    "Contract Review",
    "Tender Review",
    "Risk Assessment",
    "Meeting Minutes",
    "Progress Report",
    "Safety Inspection",
    "Business Analysis",
]
PROJECT_REVIEW_TYPES = [
    "Full Project Review",
    "Contract + Tender Review",
    "Commercial Risk Review",
    "Construction Risk Review",
    "NCC / Compliance Review",
    "Procurement Review",
    "Progress & Meeting Review",
]
EXECUTIVE_AGENT_MODES = [
    "Fast Executive Review",
    "Full Board Review",
    "Deep Due Diligence",
]
EXECUTIVE_AGENT_MODE_DESCRIPTIONS = {
    "Fast Executive Review": "Runs only the Executive Coordinator Agent. Fastest and lowest cost.",
    "Full Board Review": "Runs Contract, Risk, Finance, and Executive Coordinator agents. Balanced speed and quality.",
    "Deep Due Diligence": "Runs all executive agents. Strongest analysis and highest cost for important projects.",
}
DOCUMENT_REPORT_STRUCTURES = {
    "Contract Review": [
        "Executive Summary",
        "Key Obligations",
        "Payment Risks",
        "Variation Risks",
        "EOT / Delay Risks",
        "Dispute Clauses",
        "Red Flags",
        "Recommended Actions",
    ],
    "Tender Review": [
        "Tender Summary",
        "Scope Gaps",
        "Commercial Risks",
        "Technical Risks",
        "Clarification Questions",
        "Go / No-Go Recommendation",
    ],
    "Risk Assessment": [
        "Risk Register",
        "Safety Risks",
        "Commercial Risks",
        "Programme Risks",
        "Legal / Compliance Risks",
        "Mitigation Actions",
        "Priority Matrix",
    ],
    "Meeting Minutes": [
        "Meeting Summary",
        "Key Decisions",
        "Action Items",
        "Responsible Parties",
        "Due Dates",
        "Follow-up Items",
    ],
    "Progress Report": [
        "Executive Summary",
        "Completed Works",
        "Current Activities",
        "Delays / Issues",
        "Safety / Quality Notes",
        "Upcoming Works",
        "Action Items",
    ],
    "Safety Inspection": [
        "Hazards Identified",
        "Risk Rating",
        "Corrective Actions",
        "Responsible Person",
        "Due Date",
        "Follow-up Requirements",
    ],
    "Business Analysis": [
        "Executive Summary",
        "Key Findings",
        "Business Risks",
        "Opportunities",
        "Recommendations",
        "Next Steps",
    ],
}
DOCUMENT_PDF_FILENAMES = {
    "Contract Review": "contract_review.pdf",
    "Tender Review": "tender_review.pdf",
    "Risk Assessment": "risk_assessment.pdf",
    "Meeting Minutes": "meeting_minutes.pdf",
    "Progress Report": "progress_report.pdf",
    "Safety Inspection": "safety_inspection.pdf",
    "Business Analysis": "business_analysis.pdf",
}
PROJECT_REVIEW_PDF_FILENAMES = {
    "Full Project Review": "full_project_review.pdf",
    "Contract + Tender Review": "contract_tender_review.pdf",
    "Commercial Risk Review": "commercial_risk_review.pdf",
    "Construction Risk Review": "construction_risk_review.pdf",
    "NCC / Compliance Review": "ncc_compliance_review.pdf",
    "Procurement Review": "procurement_review.pdf",
    "Progress & Meeting Review": "progress_meeting_review.pdf",
}

INDUSTRY_OPTIONS = [
    "Construction",
    "Real Estate",
    "SaaS",
    "E-commerce",
    "Retail",
    "Manufacturing",
    "Finance",
    "Healthcare",
    "Education",
    "Hospitality",
    "Logistics",
    "Professional Services",
    "Other",
]
BUSINESS_STAGE_OPTIONS = [
    "Idea Stage",
    "Startup",
    "Growth",
    "Mature Business",
    "Expansion",
    "Turnaround",
]
STATE_OPTIONS = ["VIC", "NSW", "QLD", "WA", "SA", "TAS", "ACT", "NT"]
BUILDING_TYPE_OPTIONS = [
    "Residential",
    "Commercial",
    "Industrial",
    "Mixed-use",
    "Infrastructure",
    "Retail",
    "Education",
    "Healthcare",
    "Fit-out",
    "Renovation",
    "Extension",
    "Civil Works",
]
BUILDING_CLASS_OPTIONS = [
    "Class 1a",
    "Class 1b",
    "Class 2",
    "Class 3",
    "Class 5",
    "Class 6",
    "Class 7a",
    "Class 7b",
    "Class 8",
    "Class 9a",
    "Class 9b",
    "Class 9c",
    "Class 10a",
    "Class 10b",
    "Class 10c",
]
CONTRACT_TYPE_OPTIONS = [
    "AS4000",
    "AS4902",
    "AS2124",
    "HIA",
    "MBA",
    "Lump Sum",
    "D&C",
    "GMP",
    "Cost Plus",
    "Construction Management",
    "Bespoke Contract",
]
RISK_LEVEL_OPTIONS = ["Low", "Medium", "High", "Critical"]
STAGE_OF_WORKS_OPTIONS = [
    "Site Establishment",
    "Demolition",
    "Excavation",
    "Footings",
    "Slab",
    "Framing",
    "Lock-up",
    "Rough-in",
    "Fit-off",
    "Waterproofing",
    "Cladding",
    "Roofing",
    "Finishes",
    "Handover",
]


def field(
    name: str,
    key: str,
    field_type: str = "text_input",
    options: list[str] | None = None,
    placeholder: str = "",
) -> dict:
    return {
        "name": name,
        "key": key,
        "type": field_type,
        "options": options or [],
        "placeholder": placeholder,
    }


AUTOMATION_TOOLS = [
    "Excel",
    "Google Sheets",
    "Microsoft Project",
    "Notion",
    "Trello",
    "Asana",
    "Monday.com",
    "Slack",
    "Gmail",
    "Outlook",
    "Google Drive",
    "OneDrive",
    "Dropbox",
    "Canva",
    "Buffer",
    "Zapier",
    "Make",
    "Python",
    "Power BI",
    "Tableau",
    "Streamlit",
    "Other",
]
AUTOMATION_OUTPUT_FORMATS = [
    "Step-by-step SOP",
    "Workflow Plan",
    "Automation Blueprint",
    "Checklist",
    "Report",
    "Dashboard Plan",
    "Content Calendar",
    "Email Campaign Plan",
    "Gantt Chart Plan",
]
AUTOMATION_COMMON_FIELDS = [
    field("Business / Project Name", "business_project_name"),
    field("Industry", "industry"),
    field("Goal", "goal", "text_area"),
    field("Current Process", "current_process", "text_area"),
    field("Tools Used", "tools_used", "multiselect", AUTOMATION_TOOLS),
    field("Output Format", "output_format", "selectbox", AUTOMATION_OUTPUT_FORMATS),
]


FORM_LIBRARY = {
    "Business": [
        {
            "category": "Business",
            "name": "Market Research",
            "fields": [
                field("Industry", "industry", "selectbox", INDUSTRY_OPTIONS),
                field("Region / Country", "region_country", placeholder="e.g. Australia, New Zealand, United States"),
                field("Target Customers", "target_customers", placeholder="e.g. Residential builders, developers, contractors"),
                field("Product / Service", "product_service", placeholder="e.g. Construction project management platform"),
                field("Market Segment", "market_segment", placeholder="e.g. Residential Construction"),
                field("Business Stage", "business_stage", "selectbox", BUSINESS_STAGE_OPTIONS),
                field(
                    "Research Focus",
                    "research_focus",
                    "multiselect",
                    [
                        "Market Size",
                        "Customer Demand",
                        "Competitors",
                        "Pricing",
                        "Regulations",
                        "Technology Trends",
                        "Risks",
                        "Opportunities",
                    ],
                ),
            ],
            "template": """Conduct a professional market research report.

Industry: {industry}
Location: {region_country}
Target Customers: {target_customers}
Product / Service: {product_service}
Market Segment: {market_segment}
Business Stage: {business_stage}
Research Focus: {research_focus}

Provide:

1. Market overview
2. Market size and demand drivers
3. Customer segments
4. Competitor landscape
5. Pricing and positioning
6. Growth opportunities
7. Risks and barriers
8. Strategic recommendations""",
        },
        {
            "category": "Business",
            "name": "Product Roadmap",
            "fields": [
                field("Product", "product"),
                field("Target Users", "target_users"),
                field("Business Goal", "business_goal"),
                field("Current Stage", "current_stage", "selectbox", ["Idea", "MVP", "Beta", "Launched", "Growth", "Scaling"]),
                field("Timeline", "timeline", "selectbox", ["30 Days", "60 Days", "90 Days", "6 Months", "12 Months"]),
                field(
                    "Priority Focus",
                    "priority_focus",
                    "multiselect",
                    ["Features", "UX", "Revenue", "Automation", "Integrations", "AI Features", "Compliance", "Customer Retention"],
                ),
            ],
            "template": """Create a {timeline} product roadmap.

Product: {product}
Target Users: {target_users}
Business Goal: {business_goal}
Current Stage: {current_stage}
Priority Focus: {priority_focus}

Provide:

1. Product vision
2. Key milestones
3. Feature priorities
4. Timeline by phase
5. Resource requirements
6. Risks and dependencies
7. Success metrics
8. Recommended next actions""",
        },
        {
            "category": "Business",
            "name": "Supply Chain",
            "fields": [
                field("Business", "business"),
                field("Industry", "industry", "selectbox", INDUSTRY_OPTIONS),
                field("Key Suppliers", "key_suppliers", "text_area"),
                field("Main Products / Materials", "main_products_materials", "text_area"),
                field("Supply Chain Risk Level", "supply_chain_risk_level", "selectbox", RISK_LEVEL_OPTIONS),
                field(
                    "Procurement Focus",
                    "procurement_focus",
                    "multiselect",
                    [
                        "Cost Reduction",
                        "Supplier Diversification",
                        "Lead Time Reduction",
                        "Inventory Control",
                        "Quality Control",
                        "Sustainability",
                        "Local Sourcing",
                        "Risk Mitigation",
                    ],
                ),
            ],
            "template": """Analyze the supply chain strategy.

Business: {business}
Industry: {industry}
Key Suppliers: {key_suppliers}
Main Products / Materials: {main_products_materials}
Supply Chain Risk Level: {supply_chain_risk_level}
Procurement Focus: {procurement_focus}

Provide:

1. Procurement risks
2. Supplier dependency risks
3. Logistics and delivery risks
4. Cost reduction opportunities
5. Inventory strategy
6. Alternative supplier options
7. Process improvement plan
8. Recommendations""",
        },
        {
            "category": "Business",
            "name": "Business Analysis",
            "fields": [
                field("Business", "business"),
                field("Industry", "industry", "selectbox", INDUSTRY_OPTIONS),
                field("Current Challenge", "current_challenge", "text_area"),
                field("Business Stage", "business_stage", "selectbox", BUSINESS_STAGE_OPTIONS),
                field(
                    "Analysis Focus",
                    "analysis_focus",
                    "multiselect",
                    [
                        "SWOT",
                        "Revenue Growth",
                        "Cost Reduction",
                        "Operations",
                        "Marketing",
                        "Customer Experience",
                        "Staffing",
                        "Risk Management",
                        "Technology Adoption",
                    ],
                ),
            ],
            "template": """Perform a business analysis.

Business: {business}
Industry: {industry}
Current Challenge: {current_challenge}
Business Stage: {business_stage}
Analysis Focus: {analysis_focus}

Provide:

1. SWOT analysis
2. Revenue opportunities
3. Operational weaknesses
4. Cost improvement areas
5. Customer and market risks
6. Strategic options
7. Recommended action plan""",
        },
        {
            "category": "Business",
            "name": "Accounting",
            "fields": [
                field("Business Type", "business_type"),
                field(
                    "Revenue Range",
                    "revenue_range",
                    "selectbox",
                    ["Pre-revenue", "Under $100k", "$100k-$500k", "$500k-$1m", "$1m-$5m", "$5m+"],
                ),
                field("Expense Category", "expense_category"),
                field("Main Financial Concern", "main_financial_concern", "text_area"),
                field(
                    "Accounting Focus",
                    "accounting_focus",
                    "multiselect",
                    ["Cash Flow", "Budgeting", "Profitability", "Cost Control", "Tax Planning", "Payroll", "Reporting", "Forecasting"],
                ),
            ],
            "template": """Review the following accounting and financial scenario.

Business Type: {business_type}
Revenue: {revenue_range}
Expenses: {expense_category}
Main Financial Concern: {main_financial_concern}
Accounting Focus: {accounting_focus}

Provide:

1. Financial risk assessment
2. Cost control recommendations
3. Cash flow analysis
4. Budget improvement suggestions
5. Profitability review
6. Key financial KPIs
7. Practical next steps""",
        },
        {
            "category": "Business",
            "name": "Customer Service",
            "fields": [
                field("Business", "business"),
                field("Customer Type", "customer_type"),
                field("Current Problems", "current_problems", "text_area"),
                field(
                    "Service Channel",
                    "service_channel",
                    "multiselect",
                    ["Email", "Phone", "Live Chat", "Social Media", "Website Form", "In-person", "AI Chatbot", "CRM"],
                ),
                field(
                    "Improvement Focus",
                    "improvement_focus",
                    "multiselect",
                    [
                        "Response Time",
                        "Customer Satisfaction",
                        "Automation",
                        "Escalation Process",
                        "Staff Training",
                        "Knowledge Base",
                        "Complaint Handling",
                        "Retention",
                    ],
                ),
            ],
            "template": """Create a customer service improvement plan.

Business: {business}
Customer Type: {customer_type}
Current Problems: {current_problems}
Service Channel: {service_channel}
Improvement Focus: {improvement_focus}

Provide:

1. Key pain points
2. Response process improvements
3. Automation opportunities
4. AI chatbot opportunities
5. Customer satisfaction KPIs
6. Staff training recommendations
7. Implementation roadmap""",
        },
        {
            "category": "Business",
            "name": "Competitor Analysis",
            "fields": [
                field("Company", "company"),
                field("Industry", "industry", "selectbox", INDUSTRY_OPTIONS),
                field("Location", "location"),
                field("Main Competitors", "main_competitors", "text_area"),
                field(
                    "Comparison Focus",
                    "comparison_focus",
                    "multiselect",
                    ["Pricing", "Product Features", "Service Quality", "Branding", "Marketing", "Customer Reviews", "Technology", "Market Share"],
                ),
            ],
            "template": """Perform a competitor analysis.

Company: {company}
Industry: {industry}
Location: {location}
Main Competitors: {main_competitors}
Comparison Focus: {comparison_focus}

Provide:

1. Competitor profiles
2. Strengths and weaknesses
3. Pricing comparison
4. Market positioning
5. Differentiation opportunities
6. Threats
7. Strategic recommendations""",
        },
        {
            "category": "Business",
            "name": "Trend Analysis",
            "fields": [
                field("Industry", "industry", "selectbox", INDUSTRY_OPTIONS),
                field("Region", "region"),
                field("Time Horizon", "time_horizon", "selectbox", ["Next 3 Months", "Next 6 Months", "Next 12 Months", "2-3 Years", "5 Years"]),
                field(
                    "Trend Focus",
                    "trend_focus",
                    "multiselect",
                    [
                        "AI",
                        "Automation",
                        "Regulation",
                        "Sustainability",
                        "Customer Behavior",
                        "Economic Conditions",
                        "Technology Adoption",
                        "Labour Market",
                        "Digital Transformation",
                    ],
                ),
            ],
            "template": """Perform an industry trend analysis.

Industry: {industry}
Region: {region}
Time Horizon: {time_horizon}
Trend Focus: {trend_focus}

Provide:

1. Current market trends
2. Emerging technologies
3. Customer behavior changes
4. Regulatory or economic drivers
5. Risks
6. Opportunities
7. Strategic recommendations""",
        },
        {
            "category": "Business",
            "name": "Sales Analysis",
            "fields": [
                field("Business", "business"),
                field("Product / Service", "product_service"),
                field("Period", "period"),
                field("Sales Data / Observations", "sales_data_observations", "text_area"),
                field(
                    "Sales Channel",
                    "sales_channel",
                    "multiselect",
                    ["Direct Sales", "Website", "Marketplace", "Retail", "Wholesale", "Social Media", "Referrals", "Partnerships", "B2B", "B2C"],
                ),
                field(
                    "Sales Focus",
                    "sales_focus",
                    "multiselect",
                    ["Revenue Growth", "Conversion Rate", "Average Order Value", "Customer Acquisition", "Customer Retention", "Sales Funnel", "Pricing", "Forecasting"],
                ),
            ],
            "template": """Analyze sales performance.

Business: {business}
Product / Service: {product_service}
Period: {period}
Sales Data / Observations: {sales_data_observations}
Sales Channel: {sales_channel}
Sales Focus: {sales_focus}

Provide:

1. Revenue trend analysis
2. Customer segment analysis
3. Underperforming areas
4. Growth opportunities
5. Sales funnel improvements
6. KPI recommendations
7. Action plan""",
        },
    ],
    "Construction": [
        {
            "category": "Construction",
            "name": "Project Plan",
            "fields": [
                field("Project", "project"),
                field("State", "state", "selectbox", STATE_OPTIONS),
                field("Location", "location"),
                field("Building Type", "building_type", "selectbox", BUILDING_TYPE_OPTIONS),
                field("Building Class", "building_class", "selectbox", BUILDING_CLASS_OPTIONS),
                field("Contract Type", "contract_type", "selectbox", CONTRACT_TYPE_OPTIONS),
                field("Budget Range", "budget_range", "selectbox", ["Under $250k", "$250k-$500k", "$500k-$1m", "$1m-$5m", "$5m-$20m", "$20m+"]),
                field("Duration", "duration"),
                field("Key Constraints", "key_constraints", "text_area"),
            ],
            "template": """Create a construction project execution plan.

Project: {project}
State: {state}
Location: {location}
Building Type: {building_type}
Building Class: {building_class}
Contract Type: {contract_type}
Budget: {budget_range}
Duration: {duration}
Key Constraints: {key_constraints}

Provide:

1. Project phases
2. Scope breakdown
3. Milestones
4. Resource allocation
5. Site management approach
6. Risk management plan
7. Communication plan
8. Timeline summary""",
        },
        {
            "category": "Construction",
            "name": "Building Codes (NCC & Housing Provisions)",
            "fields": [
                field("Project", "project"),
                field("State", "state", "selectbox", STATE_OPTIONS),
                field("Building Class", "building_class", "selectbox", BUILDING_CLASS_OPTIONS),
                field("Building Type", "building_type", "selectbox", BUILDING_TYPE_OPTIONS),
                field("Scope of Works", "scope_of_works", "text_area"),
                field(
                    "Compliance Focus",
                    "compliance_focus",
                    "multiselect",
                    [
                        "Fire Safety",
                        "Structure",
                        "Waterproofing",
                        "Energy Efficiency",
                        "Accessibility",
                        "Livable Housing",
                        "Ventilation",
                        "Damp and Weatherproofing",
                        "Bushfire",
                        "Balustrades / Stairs",
                        "Wet Areas",
                        "Glazing",
                        "Roof Drainage",
                    ],
                ),
            ],
            "template": """Review this construction scenario against Australian NCC and Housing Provisions.

Project: {project}
State: {state}
Building Class: {building_class}
Building Type: {building_type}
Scope of Works: {scope_of_works}
Compliance Focus: {compliance_focus}

Provide:

1. Relevant NCC considerations
2. Housing Provision considerations if applicable
3. Compliance risks
4. Documentation required
5. Design or construction issues to check
6. Recommended actions
7. Questions to confirm with certifier / building surveyor""",
        },
        {
            "category": "Construction",
            "name": "NCC 2025 Compliance Review",
            "fields": [
                field("Project", "project"),
                field("State", "state", "selectbox", STATE_OPTIONS),
                field("Building Class", "building_class", "selectbox", BUILDING_CLASS_OPTIONS),
                field("Building Type", "building_type", "selectbox", BUILDING_TYPE_OPTIONS),
                field(
                    "Design Stage",
                    "design_stage",
                    "selectbox",
                    ["Concept Design", "Design Development", "Building Permit", "Construction Documentation", "Construction Stage", "Completion / Occupancy"],
                ),
                field("Scope", "scope", "text_area"),
                field(
                    "NCC 2025 Focus",
                    "ncc_2025_focus",
                    "multiselect",
                    ["Energy Efficiency", "Livable Housing", "Waterproofing", "Condensation", "Fire Safety", "Structure", "Accessibility", "Ventilation", "Documentation", "Performance Solution"],
                ),
            ],
            "template": """Review this project for potential NCC 2025 compliance issues.

Project: {project}
State: {state}
Building Class: {building_class}
Building Type: {building_type}
Design Stage: {design_stage}
Scope: {scope}
NCC 2025 Focus: {ncc_2025_focus}

Provide:

1. Likely NCC 2025 considerations
2. Energy efficiency considerations
3. Livable housing / accessibility considerations if applicable
4. Waterproofing, fire, structure, ventilation, and safety items to check
5. Compliance documentation checklist
6. Risk areas
7. Recommended next steps""",
        },
        {
            "category": "Construction",
            "name": "Supervision Plan",
            "fields": [
                field("Project", "project"),
                field("State", "state", "selectbox", STATE_OPTIONS),
                field("Trade / Scope", "trade_scope"),
                field("Stage of Works", "stage_of_works", "selectbox", STAGE_OF_WORKS_OPTIONS),
                field(
                    "Supervision Focus",
                    "supervision_focus",
                    "multiselect",
                    ["Quality", "Safety", "Programme", "Trade Coordination", "Defects", "Hold Points", "Site Documentation", "Subcontractor Management"],
                ),
            ],
            "template": """Create a construction site supervision plan.

Project: {project}
State: {state}
Trade / Scope: {trade_scope}
Stage of Works: {stage_of_works}
Supervision Focus: {supervision_focus}

Provide:

1. Daily supervision checklist
2. Quality control checkpoints
3. Safety responsibilities
4. Communication and reporting process
5. Inspection hold points
6. Escalation process
7. Supervisor daily report format""",
        },
        {
            "category": "Construction",
            "name": "Site Progress Report",
            "fields": [
                field("Project", "project"),
                field("Date / Week", "date_week"),
                field("Completed Works", "completed_works", "text_area"),
                field("Current Activities", "current_activities", "text_area"),
                field("Delays / Issues", "delays_issues", "text_area"),
                field("Upcoming Works", "upcoming_works", "text_area"),
                field("Progress Status", "progress_status", "selectbox", ["On Track", "Minor Delay", "Major Delay", "At Risk", "Accelerated"]),
            ],
            "template": """Generate a professional construction site progress report.

Project: {project}
Date / Week: {date_week}
Completed Works: {completed_works}
Current Activities: {current_activities}
Delays / Issues: {delays_issues}
Upcoming Works: {upcoming_works}
Progress Status: {progress_status}

Provide:

1. Executive summary
2. Completed works
3. Work in progress
4. Issues and delays
5. Safety observations
6. Quality observations
7. Upcoming activities
8. Action items""",
        },
        {
            "category": "Construction",
            "name": "Risk Assessment",
            "fields": [
                field("Project", "project"),
                field("State", "state", "selectbox", STATE_OPTIONS),
                field("Scope", "scope", "text_area"),
                field("Stage", "stage", "selectbox", STAGE_OF_WORKS_OPTIONS),
                field("Risk Level", "risk_level", "selectbox", RISK_LEVEL_OPTIONS),
                field(
                    "Risk Categories",
                    "risk_categories",
                    "multiselect",
                    ["Safety", "Cost", "Schedule", "Quality", "Environmental", "Legal / Compliance", "Design", "Procurement", "Weather", "Subcontractor Performance"],
                ),
            ],
            "template": """Perform a construction risk assessment.

Project: {project}
State: {state}
Scope: {scope}
Stage: {stage}
Risk Level: {risk_level}
Risk Categories: {risk_categories}

Identify:

1. Safety risks
2. Cost risks
3. Schedule risks
4. Quality risks
5. Environmental risks
6. Legal / compliance risks
7. Mitigation measures
8. Responsible parties""",
        },
        {
            "category": "Construction",
            "name": "Cost Estimate",
            "fields": [
                field("Project", "project"),
                field("Building Type", "building_type", "selectbox", BUILDING_TYPE_OPTIONS),
                field("Floor Area", "floor_area"),
                field("State", "state", "selectbox", STATE_OPTIONS),
                field("Scope", "scope", "text_area"),
                field(
                    "Estimate Type",
                    "estimate_type",
                    "selectbox",
                    ["Preliminary Estimate", "Concept Estimate", "Detailed Estimate", "Tender Estimate", "Variation Estimate", "Final Cost Report"],
                ),
                field(
                    "Cost Focus",
                    "cost_focus",
                    "multiselect",
                    ["Labour", "Materials", "Plant", "Subcontractors", "Preliminaries", "Contingency", "Escalation", "Risk Allowance", "GST", "Professional Fees"],
                ),
            ],
            "template": """Prepare a preliminary construction cost estimate.

Project: {project}
Building Type: {building_type}
Floor Area: {floor_area}
State: {state}
Scope: {scope}
Estimate Type: {estimate_type}
Cost Focus: {cost_focus}

Provide:

1. Major cost categories
2. Cost assumptions
3. Labour, materials, plant and subcontractor considerations
4. Contingency allowance
5. Risk allowances
6. Exclusions
7. Budget summary""",
        },
        {
            "category": "Construction",
            "name": "Safety Inspection",
            "fields": [
                field("Project", "project"),
                field("Site", "site"),
                field("Date", "date"),
                field("Observations", "observations", "text_area"),
                field(
                    "Safety Focus",
                    "safety_focus",
                    "multiselect",
                    [
                        "Working at Heights",
                        "Electrical",
                        "Excavation",
                        "Plant and Equipment",
                        "Manual Handling",
                        "Traffic Management",
                        "PPE",
                        "Housekeeping",
                        "Scaffolding",
                        "Ladders",
                        "Silica / Dust",
                        "Hot Works",
                        "Confined Space",
                    ],
                ),
                field("Risk Level", "risk_level", "selectbox", RISK_LEVEL_OPTIONS),
            ],
            "template": """Generate a construction safety inspection report.

Project: {project}
Site: {site}
Date: {date}
Observations: {observations}
Safety Focus: {safety_focus}
Risk Level: {risk_level}

Provide:

1. Hazards identified
2. Risk rating
3. Immediate corrective actions
4. Responsible persons
5. Due dates
6. Follow-up requirements
7. Safety summary""",
        },
        {
            "category": "Construction",
            "name": "Toolbox Meeting",
            "fields": [
                field("Project", "project"),
                field("Date", "date"),
                field("Topic", "topic"),
                field("Attendees", "attendees", "text_area"),
                field(
                    "Toolbox Topic",
                    "toolbox_topic",
                    "selectbox",
                    [
                        "Working at Heights",
                        "Manual Handling",
                        "Electrical Safety",
                        "Excavation Safety",
                        "Plant and Equipment",
                        "Traffic Management",
                        "PPE",
                        "Silica Dust",
                        "Housekeeping",
                        "Weather Conditions",
                        "Emergency Procedures",
                        "Site Induction",
                    ],
                ),
            ],
            "template": """Generate a toolbox meeting agenda and record.

Project: {project}
Date: {date}
Topic: {topic}
Attendees: {attendees}
Toolbox Topic: {toolbox_topic}

Provide:

1. Safety topic overview
2. Site-specific risks
3. Control measures
4. Worker responsibilities
5. Questions for discussion
6. Action items
7. Attendance record section""",
        },
        {
            "category": "Construction",
            "name": "Construction Laws Analysis",
            "fields": [
                field("Project", "project"),
                field("State", "state", "selectbox", STATE_OPTIONS),
                field("Contract Type", "contract_type", "selectbox", CONTRACT_TYPE_OPTIONS),
                field("Issue", "issue", "text_area"),
                field(
                    "Legal Focus",
                    "legal_focus",
                    "multiselect",
                    ["Payment Claims", "Security of Payment", "Variations", "EOT", "Defects", "Latent Conditions", "Liquidated Damages", "Termination", "WHS", "Licensing", "Insurance", "Dispute Resolution"],
                ),
            ],
            "template": """Analyze applicable construction laws and regulatory risks.

Project: {project}
State: {state}
Contract Type: {contract_type}
Issue: {issue}
Legal Focus: {legal_focus}

Provide:

1. Relevant legal and regulatory considerations
2. Contractual obligations
3. Compliance risks
4. Documentation required
5. Possible consequences
6. Recommended practical actions
7. Items requiring legal advice""",
        },
        {
            "category": "Construction",
            "name": "Tender Review",
            "fields": [
                field("Project", "project"),
                field("Client", "client"),
                field("Contract Type", "contract_type", "selectbox", CONTRACT_TYPE_OPTIONS),
                field("Tender Scope", "tender_scope", "text_area"),
                field("Tender Value", "tender_value"),
                field(
                    "Review Focus",
                    "review_focus",
                    "multiselect",
                    ["Commercial Risk", "Technical Risk", "Scope Gap", "Programme Risk", "Liquidated Damages", "Payment Terms", "Design Responsibility", "Latent Conditions", "Exclusions", "Clarifications"],
                ),
            ],
            "template": """Review the following construction tender submission.

Project: {project}
Client: {client}
Contract Type: {contract_type}
Tender Scope: {tender_scope}
Tender Value: {tender_value}
Review Focus: {review_focus}

Provide:

1. Commercial risks
2. Technical risks
3. Scope gaps
4. Ambiguous clauses
5. Missing documentation
6. Pricing risks
7. Clarification questions
8. Recommendation: proceed / proceed with conditions / do not proceed""",
        },
        {
            "category": "Construction",
            "name": "RFI Generator",
            "fields": [
                field("Project", "project"),
                field("Drawing / Specification Reference", "drawing_specification_reference"),
                field("Issue", "issue", "text_area"),
                field("Impact", "impact", "text_area"),
                field(
                    "RFI Type",
                    "rfi_type",
                    "selectbox",
                    ["Design Clarification", "Specification Conflict", "Drawing Discrepancy", "Site Condition", "Material Substitution", "Coordination Issue", "Programme Impact", "Cost Impact"],
                ),
            ],
            "template": """Create a professional Request for Information.

Project: {project}
Drawing / Specification Reference: {drawing_specification_reference}
Issue: {issue}
Impact: {impact}
RFI Type: {rfi_type}

Provide:

1. RFI title
2. Background
3. Specific question
4. Reason for request
5. Potential cost / time impact
6. Required response date
7. Professional wording""",
        },
        {
            "category": "Construction",
            "name": "Variation Claim",
            "fields": [
                field("Project", "project"),
                field("Contract Type", "contract_type", "selectbox", CONTRACT_TYPE_OPTIONS),
                field("Variation Description", "variation_description", "text_area"),
                field("Cause", "cause", "text_area"),
                field("Cost Impact", "cost_impact"),
                field("Time Impact", "time_impact"),
                field(
                    "Variation Type",
                    "variation_type",
                    "selectbox",
                    ["Client Instruction", "Design Change", "Latent Condition", "Scope Gap", "Regulatory Requirement", "Site Condition", "Delay Event", "Material Substitution"],
                ),
            ],
            "template": """Prepare a construction variation claim.

Project: {project}
Contract: {contract_type}
Variation Description: {variation_description}
Cause: {cause}
Cost Impact: {cost_impact}
Time Impact: {time_impact}
Variation Type: {variation_type}

Provide:

1. Variation summary
2. Contract basis
3. Scope change explanation
4. Cost breakdown
5. Time impact
6. Supporting evidence checklist
7. Formal claim wording""",
        },
        {
            "category": "Construction",
            "name": "Subcontractor Evaluation",
            "fields": [
                field("Subcontractor", "subcontractor"),
                field(
                    "Trade",
                    "trade",
                    "selectbox",
                    ["Demolition", "Excavation", "Concrete", "Steel", "Carpentry", "Roofing", "Cladding", "Waterproofing", "Plumbing", "Electrical", "HVAC", "Fire Services", "Plastering", "Painting", "Tiling", "Flooring", "Landscaping", "Joinery", "Glazing"],
                ),
                field("Project", "project"),
                field("Performance Observations", "performance_observations", "text_area"),
                field(
                    "Evaluation Focus",
                    "evaluation_focus",
                    "multiselect",
                    ["Quality", "Safety", "Programme", "Communication", "Commercial", "Defects", "Documentation", "Coordination", "Responsiveness"],
                ),
            ],
            "template": """Evaluate a subcontractor.

Subcontractor: {subcontractor}
Trade: {trade}
Project: {project}
Performance Observations: {performance_observations}
Evaluation Focus: {evaluation_focus}

Provide:

1. Capability assessment
2. Quality performance
3. Safety performance
4. Programme performance
5. Communication performance
6. Commercial risks
7. Recommendation
8. Scorecard format""",
        },
        {
            "category": "Construction",
            "name": "Contract Analysis",
            "fields": [
                field("Project", "project"),
                field("Contract Type", "contract_type", "selectbox", CONTRACT_TYPE_OPTIONS),
                field("Clause / Issue", "clause_issue", "text_area"),
                field("Concern", "concern", "text_area"),
                field(
                    "Contract Focus",
                    "contract_focus",
                    "multiselect",
                    ["Key Obligations", "Risk Clauses", "Payment Terms", "Variations", "EOT", "Delay Damages", "Defects", "Dispute Resolution", "Termination", "Insurance", "Warranties", "Scope Exclusions"],
                ),
            ],
            "template": """Analyze this construction contract.

Project: {project}
Contract Type: {contract_type}
Clause / Issue: {clause_issue}
Concern: {concern}
Contract Focus: {contract_focus}

Provide:

1. Key obligations
2. Risk clauses
3. Payment terms
4. Variation process
5. Delay and extension of time clauses
6. Dispute resolution process
7. Practical recommendations
8. Items requiring legal review""",
        },
        {
            "category": "Construction",
            "name": "BIM + AI",
            "fields": [
                field("Project", "project"),
                field("Stage", "stage", "selectbox", STAGE_OF_WORKS_OPTIONS),
                field(
                    "Current BIM Platform",
                    "current_bim_platform",
                    "selectbox",
                    ["Revit", "Navisworks", "ArchiCAD", "Tekla", "BIM360", "Autodesk Construction Cloud", "Procore", "Bluebeam", "Other"],
                ),
                field("Main Problems", "main_problems", "text_area"),
                field(
                    "BIM / AI Focus",
                    "bim_ai_focus",
                    "multiselect",
                    ["Clash Detection", "Quantity Takeoff", "Programme Analysis", "Cost Control", "Site Reporting", "Defect Detection", "Document Control", "Safety Analytics", "Digital Twin", "Procurement Automation"],
                ),
            ],
            "template": """Analyze BIM and AI opportunities for this construction project.

Project: {project}
Stage: {stage}
Current BIM Platform: {current_bim_platform}
Main Problems: {main_problems}
BIM / AI Focus: {bim_ai_focus}

Provide:

1. BIM workflow improvements
2. AI automation opportunities
3. Clash detection benefits
4. Quantity takeoff opportunities
5. Programme and cost control opportunities
6. Site reporting automation
7. Implementation roadmap
8. Risks and limitations""",
        },
        {
            "category": "Construction",
            "name": "Digital Twin Strategy",
            "fields": [
                field("Project", "project"),
                field(
                    "Asset Type",
                    "asset_type",
                    "selectbox",
                    ["Residential Building", "Commercial Building", "Industrial Facility", "Hospital", "School", "Retail Centre", "Infrastructure Asset", "Civil Asset", "Mixed-use Development"],
                ),
                field("Project Stage", "project_stage", "selectbox", STAGE_OF_WORKS_OPTIONS),
                field("Available Data", "available_data", "text_area"),
                field(
                    "Digital Twin Focus",
                    "digital_twin_focus",
                    "multiselect",
                    ["Asset Management", "Maintenance", "Energy Performance", "Occupancy", "Safety", "IoT Sensors", "BIM Integration", "Operations", "Lifecycle Cost"],
                ),
            ],
            "template": """Create a digital twin strategy for a construction or asset project.

Project: {project}
Asset Type: {asset_type}
Project Stage: {project_stage}
Available Data: {available_data}
Digital Twin Focus: {digital_twin_focus}

Provide:

1. Digital twin objectives
2. Required data sources
3. BIM / IoT / sensor integration
4. Operations and maintenance use cases
5. Cost and performance benefits
6. Implementation roadmap
7. Risks and governance requirements""",
        },
        {
            "category": "Construction",
            "name": "Prefabrication / Modular Construction",
            "fields": [
                field("Project", "project"),
                field("Building Type", "building_type", "selectbox", BUILDING_TYPE_OPTIONS),
                field("State", "state", "selectbox", STATE_OPTIONS),
                field("Scope", "scope", "text_area"),
                field(
                    "Prefab Focus",
                    "prefab_focus",
                    "multiselect",
                    ["Bathroom Pods", "Wall Panels", "Roof Trusses", "Structural Steel", "Precast Concrete", "Modular Rooms", "Services Risers", "Facade Panels", "Programme Reduction", "Waste Reduction"],
                ),
            ],
            "template": """Assess prefabrication or modular construction opportunities.

Project: {project}
Building Type: {building_type}
State: {state}
Scope: {scope}
Prefab Focus: {prefab_focus}

Provide:

1. Suitable prefabrication elements
2. Benefits
3. Cost and programme impacts
4. Design coordination requirements
5. Logistics considerations
6. Risks
7. Recommendation""",
        },
        {
            "category": "Construction",
            "name": "Construction Productivity Improvement",
            "fields": [
                field("Project", "project"),
                field("Trade / Scope", "trade_scope"),
                field("Current Bottlenecks", "current_bottlenecks", "text_area"),
                field(
                    "Productivity Focus",
                    "productivity_focus",
                    "multiselect",
                    ["Labour Utilisation", "Material Handling", "Trade Sequencing", "Site Logistics", "Rework Reduction", "Digital Tools", "AI Automation", "Standardisation", "Lean Construction"],
                ),
            ],
            "template": """Create a construction productivity improvement plan.

Project: {project}
Trade / Scope: {trade_scope}
Current Bottlenecks: {current_bottlenecks}
Productivity Focus: {productivity_focus}

Provide:

1. Productivity issues
2. Workflow improvements
3. Labour and resource optimisation
4. Digital tool opportunities
5. AI automation opportunities
6. KPIs to track
7. Implementation plan""",
        },
        {
            "category": "Construction",
            "name": "Defect / Quality Inspection",
            "fields": [
                field("Project", "project"),
                field("Area", "area"),
                field("Inspection Date", "inspection_date"),
                field("Observed Defects", "observed_defects", "text_area"),
                field(
                    "Defect Category",
                    "defect_category",
                    "selectbox",
                    ["Waterproofing", "Structural", "Finishes", "Plumbing", "Electrical", "Fire Services", "Doors / Windows", "Cladding", "Roofing", "Painting", "Tiling", "Flooring", "Joinery", "Documentation"],
                ),
                field("Severity", "severity", "selectbox", RISK_LEVEL_OPTIONS),
            ],
            "template": """Generate a defect and quality inspection report.

Project: {project}
Area: {area}
Inspection Date: {inspection_date}
Observed Defects: {observed_defects}
Defect Category: {defect_category}
Severity: {severity}

Provide:

1. Defect list
2. Severity rating
3. Likely causes
4. Required rectification
5. Responsible party
6. Due date
7. Reinspection checklist""",
        },
        {
            "category": "Construction",
            "name": "Procurement Risk Review",
            "fields": [
                field("Project", "project"),
                field("Key Materials", "key_materials", "text_area"),
                field("Suppliers", "suppliers", "text_area"),
                field("Programme Requirements", "programme_requirements", "text_area"),
                field(
                    "Procurement Focus",
                    "procurement_focus",
                    "multiselect",
                    ["Long-lead Items", "Supplier Risk", "Price Escalation", "Logistics", "Substitution Options", "Local Availability", "Quality Risk", "Storage", "Delivery Sequencing", "Contract Terms"],
                ),
            ],
            "template": """Review procurement risks for a construction project.

Project: {project}
Key Materials: {key_materials}
Suppliers: {suppliers}
Programme Requirements: {programme_requirements}
Procurement Focus: {procurement_focus}

Provide:

1. Long-lead items
2. Supplier risks
3. Price escalation risks
4. Logistics risks
5. Alternative procurement options
6. Mitigation strategy
7. Procurement action plan""",
        },
    ],
}


AUTOMATION_LIBRARY = [
    {
        "name": "Project Management Automation",
        "fields": [
            *AUTOMATION_COMMON_FIELDS,
            field("Project Type", "project_type", "selectbox", ["Construction Project", "Software Project", "Marketing Project", "Business Operations", "Product Launch", "Internal Workflow", "Other"]),
            field("Team Size", "team_size"),
            field("Reporting Frequency", "reporting_frequency", "selectbox", ["Daily", "Weekly", "Fortnightly", "Monthly"]),
            field("Current Project Management Tool", "current_project_management_tool"),
            field("Pain Points", "pain_points", "text_area"),
            field(
                "Automation Scope",
                "automation_scope",
                "multiselect",
                ["Task Assignment", "Status Reporting", "Risk Register", "Issue Tracking", "Meeting Notes", "Progress Dashboard", "Reminder Automation", "Document Control"],
            ),
        ],
        "output_sections": [
            "Executive Summary",
            "Current PM Workflow Assessment",
            "Automated Task Breakdown",
            "RACI Responsibility Matrix",
            "Risk / Issue Tracking Workflow",
            "Weekly Reporting Automation",
            "Tool Setup Recommendation",
            "Implementation Checklist",
            "KPI Metrics",
        ],
        "template": """Create a professional project management automation blueprint.

Business / Project Name: {business_project_name}
Industry: {industry}
Goal: {goal}
Current Process: {current_process}
Tools Used: {tools_used}
Preferred Output Format: {output_format}

Project Type: {project_type}
Team Size: {team_size}
Reporting Frequency: {reporting_frequency}
Current Project Management Tool: {current_project_management_tool}
Pain Points: {pain_points}
Automation Scope: {automation_scope}

Design a practical automation system that improves project control, reporting, accountability, and follow-up without requiring unsupported API access.

Provide:
{output_sections}

Include this table:
Task | Owner | Frequency | Trigger | Automation Tool | Output""",
    },
    {
        "name": "Gantt Chart Planner",
        "fields": [
            *AUTOMATION_COMMON_FIELDS,
            field("Project Start Date", "project_start_date"),
            field("Project End Date", "project_end_date"),
            field("Major Phases", "major_phases", "text_area"),
            field("Key Milestones", "key_milestones", "text_area"),
            field("Resource Constraints", "resource_constraints", "text_area"),
            field("Dependency Notes", "dependency_notes", "text_area"),
        ],
        "output_sections": [
            "Work Breakdown Structure",
            "Phase Plan",
            "Milestone Plan",
            "Dependency Map",
            "Critical Path Assumptions",
            "Resource Allocation",
            "Delay Monitoring Plan",
            "Gantt Chart Table",
        ],
        "template": """Create a professional Gantt chart planning prompt and schedule blueprint.

Business / Project Name: {business_project_name}
Industry: {industry}
Goal: {goal}
Current Process: {current_process}
Tools Used: {tools_used}
Preferred Output Format: {output_format}

Project Start Date: {project_start_date}
Project End Date: {project_end_date}
Major Phases: {major_phases}
Key Milestones: {key_milestones}
Resource Constraints: {resource_constraints}
Dependency Notes: {dependency_notes}

Build a realistic schedule plan with phases, dependencies, owners, and monitoring controls.

Provide:
{output_sections}

Include this table:
Task | Phase | Start Date | End Date | Duration | Dependency | Owner | Status""",
    },
    {
        "name": "Social Media Auto Posting Planner",
        "fields": [
            *AUTOMATION_COMMON_FIELDS,
            field("Brand Name", "brand_name"),
            field("Target Audience", "target_audience", "text_area"),
            field("Platforms", "platforms", "multiselect", ["LinkedIn", "Instagram", "Facebook", "X / Twitter", "TikTok", "YouTube Shorts"]),
            field("Posting Frequency", "posting_frequency", "selectbox", ["Daily", "3 times per week", "Weekly", "Fortnightly", "Monthly"]),
            field("Content Pillars", "content_pillars", "text_area"),
            field("Tone", "tone", "selectbox", ["Professional", "Educational", "Friendly", "Sales-oriented", "Thought Leadership", "Luxury / Premium", "Technical"]),
            field("Approval Requirement", "approval_requirement"),
            field("Asset Source", "asset_source", "selectbox", ["Original photos", "AI-generated images", "Licensed stock images", "User-provided assets", "No images"]),
        ],
        "output_sections": [
            "Platform Strategy",
            "30-Day Content Calendar",
            "Post Ideas",
            "Captions",
            "Hashtags",
            "Image / Video Asset Suggestions",
            "Approval Workflow",
            "Copyright-Safe Content Rules",
            "Scheduling Workflow",
        ],
        "template": """Create a social media scheduling and approval automation plan.

Business / Project Name: {business_project_name}
Industry: {industry}
Goal: {goal}
Current Process: {current_process}
Tools Used: {tools_used}
Preferred Output Format: {output_format}

Brand Name: {brand_name}
Target Audience: {target_audience}
Platforms: {platforms}
Posting Frequency: {posting_frequency}
Content Pillars: {content_pillars}
Tone: {tone}
Approval Requirement: {approval_requirement}
Asset Source: {asset_source}

Create a practical content planning workflow for human-reviewed scheduling. Do not create instructions to auto-publish without review.

Provide:
{output_sections}

Include this table:
Date | Platform | Topic | Caption | Asset Suggestion | Hashtags | Status

Safety rules:
- Do not use copyrighted images without permission
- Do not use brand logos unless authorised
- Prefer original, AI-generated, licensed, or user-provided assets
- Human approval required before publishing
- Do not auto-publish without review""",
    },
    {
        "name": "Market Research Automation",
        "fields": [
            *AUTOMATION_COMMON_FIELDS,
            field("Target Market", "target_market"),
            field("Competitors", "competitors", "text_area"),
            field("Research Frequency", "research_frequency", "selectbox", ["Weekly", "Monthly", "Quarterly"]),
            field(
                "Data Sources",
                "data_sources",
                "multiselect",
                ["Company Websites", "Google Search", "LinkedIn", "Industry Reports", "Government Data", "News", "Customer Reviews", "Social Media", "Internal Sales Data"],
            ),
            field("KPIs to Track", "kpis_to_track", "text_area"),
            field("Report Audience", "report_audience"),
        ],
        "output_sections": [
            "Market Research Automation Workflow",
            "Competitor Tracking System",
            "Trend Monitoring System",
            "Source List",
            "KPI Dashboard Plan",
            "Alert Rules",
            "Reporting Schedule",
            "Monthly Report Template",
        ],
        "template": """Create a market research automation workflow and reporting system.

Business / Project Name: {business_project_name}
Industry: {industry}
Goal: {goal}
Current Process: {current_process}
Tools Used: {tools_used}
Preferred Output Format: {output_format}

Target Market: {target_market}
Competitors: {competitors}
Research Frequency: {research_frequency}
Data Sources: {data_sources}
KPIs to Track: {kpis_to_track}
Report Audience: {report_audience}

Design a repeatable research process for monitoring competitors, market trends, and KPIs using approved sources and human review.

Provide:
{output_sections}

Include this table:
Data Source | What to Track | Frequency | Tool | Output""",
    },
    {
        "name": "Repetitive Task Automation",
        "fields": [
            *AUTOMATION_COMMON_FIELDS,
            field("Repetitive Task Name", "repetitive_task_name"),
            field("Manual Steps", "manual_steps", "text_area"),
            field("Frequency", "frequency", "selectbox", ["Multiple times per day", "Daily", "Weekly", "Monthly"]),
            field("Time Spent", "time_spent"),
            field("Error Risks", "error_risks", "text_area"),
            field("Desired Outcome", "desired_outcome", "text_area"),
        ],
        "output_sections": [
            "Task Mapping",
            "Manual Step Analysis",
            "Automation Opportunities",
            "Trigger / Action Workflow",
            "Tool Recommendation",
            "SOP",
            "Risk Controls",
            "Time Saving Estimate",
        ],
        "template": """Create a repetitive task automation blueprint.

Business / Project Name: {business_project_name}
Industry: {industry}
Goal: {goal}
Current Process: {current_process}
Tools Used: {tools_used}
Preferred Output Format: {output_format}

Repetitive Task Name: {repetitive_task_name}
Manual Steps: {manual_steps}
Frequency: {frequency}
Time Spent: {time_spent}
Error Risks: {error_risks}
Desired Outcome: {desired_outcome}

Map the manual task, identify automation opportunities, and create a practical SOP with controls.

Provide:
{output_sections}

Include this table:
Manual Step | Automation Opportunity | Trigger | Action | Tool | Risk Control""",
    },
    {
        "name": "Data Analysis & Visualization",
        "fields": [
            *AUTOMATION_COMMON_FIELDS,
            field("Data Source", "data_source"),
            field("Data Format", "data_format", "selectbox", ["CSV", "Excel", "Google Sheets", "SQL Database", "API", "Manual Entry"]),
            field("Metrics / KPIs", "metrics_kpis", "text_area"),
            field("Audience", "audience"),
            field("Dashboard Tool", "dashboard_tool", "selectbox", ["Excel", "Google Sheets", "Power BI", "Tableau", "Looker Studio", "Streamlit", "Python"]),
            field("Chart Types", "chart_types", "multiselect", ["Bar Chart", "Line Chart", "Pie Chart", "Scatter Plot", "Heatmap", "Gantt Chart", "KPI Cards", "Dashboard", "Funnel Chart"]),
            field("Reporting Frequency", "reporting_frequency"),
        ],
        "output_sections": [
            "Data Pipeline",
            "Data Cleaning Steps",
            "KPI Definitions",
            "Dashboard Layout",
            "Recommended Charts",
            "Automated Reporting Schedule",
            "Data Quality Checks",
            "Executive Insight Summary",
        ],
        "template": """Create a data analysis and visualization automation plan.

Business / Project Name: {business_project_name}
Industry: {industry}
Goal: {goal}
Current Process: {current_process}
Tools Used: {tools_used}
Preferred Output Format: {output_format}

Data Source: {data_source}
Data Format: {data_format}
Metrics / KPIs: {metrics_kpis}
Audience: {audience}
Dashboard Tool: {dashboard_tool}
Chart Types: {chart_types}
Reporting Frequency: {reporting_frequency}

Design a practical analytics workflow with clean data inputs, KPI logic, visualization choices, and reporting cadence.

Provide:
{output_sections}

Include this table:
KPI | Data Source | Calculation | Chart Type | Update Frequency""",
    },
    {
        "name": "Office Automation",
        "fields": [
            *AUTOMATION_COMMON_FIELDS,
            field("Department", "department", "selectbox", ["Admin", "Sales", "Operations", "Construction", "Finance", "HR", "Marketing", "Customer Service"]),
            field("Admin Tasks", "admin_tasks", "text_area"),
            field("Documents Used", "documents_used", "text_area"),
            field("Approval Process", "approval_process", "text_area"),
            field("Communication Channels", "communication_channels"),
            field("Bottlenecks", "bottlenecks", "text_area"),
        ],
        "output_sections": [
            "Office Workflow Assessment",
            "Document Template System",
            "Email Automation Plan",
            "Calendar Automation Plan",
            "Approval Workflow",
            "Filing System",
            "Reporting Routine",
            "SOP",
        ],
        "template": """Create an office automation plan for a department workflow.

Business / Project Name: {business_project_name}
Industry: {industry}
Goal: {goal}
Current Process: {current_process}
Tools Used: {tools_used}
Preferred Output Format: {output_format}

Department: {department}
Admin Tasks: {admin_tasks}
Documents Used: {documents_used}
Approval Process: {approval_process}
Communication Channels: {communication_channels}
Bottlenecks: {bottlenecks}

Design a practical office workflow with templates, approvals, filing, communication, and reporting routines.

Provide:
{output_sections}

Include this table:
Office Task | Current Process | Automation Method | Tool | Owner | Frequency""",
    },
    {
        "name": "File Organization Automation",
        "fields": [
            *AUTOMATION_COMMON_FIELDS,
            field("File Types", "file_types", "text_area"),
            field("Current Folder Structure", "current_folder_structure", "text_area"),
            field("Naming Problems", "naming_problems", "text_area"),
            field("Storage Platform", "storage_platform", "selectbox", ["Google Drive", "OneDrive", "Dropbox", "Local Server", "SharePoint", "Other"]),
            field("Access Rules", "access_rules", "text_area"),
            field("Retention Requirements", "retention_requirements", "text_area"),
        ],
        "output_sections": [
            "Folder Structure",
            "Naming Convention",
            "Auto-Tagging Rules",
            "File Routing Workflow",
            "Backup Process",
            "Version Control Rules",
            "Archive Policy",
            "Access Permission Matrix",
        ],
        "template": """Create a file organization automation blueprint.

Business / Project Name: {business_project_name}
Industry: {industry}
Goal: {goal}
Current Process: {current_process}
Tools Used: {tools_used}
Preferred Output Format: {output_format}

File Types: {file_types}
Current Folder Structure: {current_folder_structure}
Naming Problems: {naming_problems}
Storage Platform: {storage_platform}
Access Rules: {access_rules}
Retention Requirements: {retention_requirements}

Design a controlled filing, naming, access, backup, and retention system that can be implemented with the selected storage platform.

Provide:
{output_sections}

Include this table:
File Type | Folder Path | Naming Rule | Owner | Retention Period | Access Level""",
    },
    {
        "name": "Promotion Email Automation",
        "fields": [
            *AUTOMATION_COMMON_FIELDS,
            field("Campaign Name", "campaign_name"),
            field("Product / Service", "product_service"),
            field("Target Audience", "target_audience", "text_area"),
            field("Offer", "offer", "text_area"),
            field("Email Sequence Length", "email_sequence_length", "selectbox", ["3 emails", "5 emails", "7 emails"]),
            field("Tone", "tone"),
            field("CTA", "cta"),
            field("Compliance Region", "compliance_region", "selectbox", ["Australia", "United States", "United Kingdom", "European Union", "Other"]),
        ],
        "output_sections": [
            "Campaign Strategy",
            "Audience Segments",
            "Email Sequence",
            "Subject Lines",
            "Email Body Drafts",
            "CTA Recommendations",
            "Follow-up Schedule",
            "Compliance Checklist",
            "Performance Metrics",
        ],
        "template": """Create a promotion email automation campaign plan.

Business / Project Name: {business_project_name}
Industry: {industry}
Goal: {goal}
Current Process: {current_process}
Tools Used: {tools_used}
Preferred Output Format: {output_format}

Campaign Name: {campaign_name}
Product / Service: {product_service}
Target Audience: {target_audience}
Offer: {offer}
Email Sequence Length: {email_sequence_length}
Tone: {tone}
CTA: {cta}
Compliance Region: {compliance_region}

Create a compliant, human-reviewed email campaign plan. Do not send emails, scrape contacts, or use purchased lists.

Provide:
{output_sections}

Include this table:
Email # | Timing | Subject Line | Goal | CTA | Follow-up Condition

Compliance reminders:
- Include unsubscribe option
- Avoid misleading subject lines
- Respect consent and privacy
- Do not scrape or use purchased emails without permission
- Follow applicable spam and privacy laws""",
    },
]


def run_fast_mode(query: str) -> dict:
    from crewai import Crew, Task

    import intelligent_gateway

    task = Task(
        description=query,
        expected_output="A professional detailed response.",
        agent=intelligent_gateway.writer_agent,
    )
    crew = Crew(
        agents=[intelligent_gateway.writer_agent],
        tasks=[task],
        verbose=True,
    )

    try:
        result = crew.kickoff()
    except Exception:
        return {
            "category": "writing",
            "confidence": 1.0,
            "result": "Gateway execution failed: unable to reach the LLM provider. Please try again later.",
            "version": VERSION,
        }

    try:
        output = result.raw
    except AttributeError:
        output = str(result)

    return {
        "category": "writing",
        "confidence": 1.0,
        "result": output,
        "version": VERSION,
    }


def run_advanced_mode(query: str) -> dict:
    response = requests.post(
        API_URL,
        json={"query": query},
        timeout=REQUEST_TIMEOUT_SECONDS,
    )
    response.raise_for_status()
    return response.json()


def save_history_entry(
    query: str,
    mode_name: str,
    payload: dict,
    elapsed_seconds: float,
    metadata: dict | None = None,
) -> dict:
    entry = {
        "query": query,
        "mode": mode_name,
        "category": payload.get("category", ""),
        "confidence": payload.get("confidence", ""),
        "version": payload.get("version", ""),
        "result": payload.get("result", ""),
        "elapsed_seconds": elapsed_seconds,
    }
    if metadata:
        entry.update(metadata)
    if payload.get("agent_mode"):
        entry["agent_mode"] = payload["agent_mode"]
    st.session_state.history.insert(0, entry)
    st.session_state.history = st.session_state.history[:HISTORY_LIMIT]
    st.session_state.selected_history = entry
    return entry


def get_document_analysis_type(query: str) -> str | None:
    for line in query.splitlines():
        if line.startswith("Analysis Type:"):
            analysis_type = line.split(":", 1)[1].strip()
            if analysis_type in DOCUMENT_PDF_FILENAMES:
                return analysis_type
    return None


def get_project_review_type(query: str) -> str | None:
    for line in query.splitlines():
        if line.startswith("Review Type:"):
            review_type = line.split(":", 1)[1].strip()
            if review_type in PROJECT_REVIEW_PDF_FILENAMES:
                return review_type
    return None


def build_pdf_report(
    analysis_type: str,
    report_content: str,
    agent_mode: str | None = None,
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    generated_time = time.strftime("%Y-%m-%d %H:%M:%S")
    buffer = BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
    )
    base_styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=base_styles["Title"],
        fontName="STSong-Light",
        fontSize=18,
        leading=24,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "ReportMeta",
        parent=base_styles["Normal"],
        fontName="STSong-Light",
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    content_style = ParagraphStyle(
        "ReportContent",
        parent=base_styles["Normal"],
        fontName="STSong-Light",
        fontSize=10,
        leading=14,
        spaceAfter=4,
    )
    heading_style = ParagraphStyle(
        "ReportHeading",
        parent=base_styles["Heading2"],
        fontName="STSong-Light",
        fontSize=12,
        leading=16,
        spaceBefore=8,
        spaceAfter=6,
    )

    story = [
        Paragraph("Gao Intelligence Hub", title_style),
        Paragraph(f"Analysis Type: {escape(analysis_type)}", meta_style),
        Paragraph(f"Generated Time: {escape(generated_time)}", meta_style),
    ]
    if agent_mode:
        story.append(Paragraph(f"Agent Mode: {escape(agent_mode)}", meta_style))
    story.extend([Spacer(1, 8), Paragraph("Report Content", heading_style)])

    for line in report_content.splitlines():
        if line.strip():
            story.append(Paragraph(escape(line), content_style))
        else:
            story.append(Spacer(1, 6))

    document.build(story)
    return buffer.getvalue()


def render_pdf_download(entry: dict) -> None:
    analysis_type = get_document_analysis_type(entry.get("query", ""))
    review_type = get_project_review_type(entry.get("query", ""))
    report_content = str(entry.get("result", ""))
    report_type = analysis_type or review_type
    if not report_type or not report_content:
        return

    pdf_bytes = build_pdf_report(
        report_type,
        report_content,
        entry.get("agent_mode"),
    )
    file_name = DOCUMENT_PDF_FILENAMES.get(
        report_type,
        PROJECT_REVIEW_PDF_FILENAMES.get(report_type, "executive_review.pdf"),
    )
    st.download_button(
        "Download Report PDF",
        data=pdf_bytes,
        file_name=file_name,
        mime="application/pdf",
        use_container_width=True,
    )


def inject_custom_css() -> None:
    st.markdown(
        """
        <style>
        :root {
            --hub-bg: #F8FAFC;
            --hub-card: #FFFFFF;
            --hub-border: #E2E8F0;
            --hub-text: #0F172A;
            --hub-muted: #475569;
            --hub-accent: #2563EB;
            --hub-purple: #7C3AED;
            --hub-success: #16A34A;
            --hub-warning: #EA580C;
            --hub-accent-soft: #EFF6FF;
        }

        html,
        body {
            width: 100%;
            overflow-x: hidden !important;
        }

        *,
        *::before,
        *::after {
            box-sizing: border-box;
        }

        .stApp {
            background: var(--hub-bg);
            color: var(--hub-text) !important;
            overflow-x: hidden !important;
            width: 100%;
        }

        .main .block-container,
        .block-container {
            max-width: 100% !important;
            padding-top: 2rem !important;
            padding-left: 3rem !important;
            padding-right: 3rem !important;
            padding-bottom: 3rem !important;
            width: 100% !important;
        }

        h1, h2, h3 {
            color: var(--hub-text) !important;
            letter-spacing: 0;
        }

        h1 {
            font-size: clamp(24px, 3rem, 48px);
        }

        h2 {
            font-size: clamp(18px, 2rem, 32px);
        }

        label, p, span, div {
            color: #0F172A;
        }

        p {
            color: #475569 !important;
            font-size: clamp(14px, 1rem, 16px);
        }

        p, span, div, small {
            letter-spacing: 0;
        }

        img {
            max-width: 100%;
            height: auto;
        }

        input,
        textarea,
        select {
            max-width: 100% !important;
            width: 100% !important;
        }

        button {
            border-radius: 12px !important;
            min-height: 48px;
            max-width: 100%;
            width: 100%;
        }

        .table-wrapper,
        [data-testid="stTable"],
        [data-testid="stDataFrame"] {
            max-width: 100%;
            overflow-x: auto;
        }

        .grid {
            display: grid;
            gap: 16px;
            width: 100%;
            max-width: 100%;
        }

        .card,
        .section,
        .hero,
        .hub-hero,
        .status-card-native,
        .mobile-guide-card,
        .hub-result-body,
        .subscription-mobile-card,
        .subscription-desktop-metrics {
            width: 100%;
            max-width: 100%;
        }

        @media (min-width: 1200px) {
            .grid {
                grid-template-columns: repeat(4, minmax(0, 1fr));
            }
        }

        @media (max-width: 1199px) {
            .grid {
                grid-template-columns: repeat(2, minmax(0, 1fr));
            }
        }

        @media (max-width: 768px) {
            .grid {
                grid-template-columns: 1fr;
            }
        }

        [data-testid="InputInstructions"],
        div[class*="InputInstructions"],
        div[class*="inputInstructions"],
        div[class*="InputInstructions"] *,
        div[class*="inputInstructions"] * {
            display: none !important;
            height: 0 !important;
            min-height: 0 !important;
            margin: 0 !important;
            padding: 0 !important;
            visibility: hidden !important;
        }

        [data-testid="stSidebar"] {
            background: #0f172a !important;
            color: white !important;
        }

        [data-testid="stSidebar"] * {
            color: white !important;
        }

        section[data-testid="stSidebar"] {
            background: #0f172a !important;
            border-right: 1px solid rgba(148, 163, 184, 0.18);
        }

        section[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] h2 {
            color: #FFFFFF !important;
            font-size: 1.2rem;
            font-weight: 800;
            line-height: 1.2;
            margin-bottom: 1rem;
        }

        section[data-testid="stSidebar"] [data-testid="stCaptionContainer"],
        section[data-testid="stSidebar"] [data-testid="stCaptionContainer"] * {
            color: #CBD5E1 !important;
        }

        section[data-testid="stSidebar"] [data-testid="stRadio"] label,
        section[data-testid="stSidebar"] [data-testid="stRadio"] label *,
        section[data-testid="stSidebar"] [role="radiogroup"] *,
        section[data-testid="stSidebar"] [data-testid="stWidgetLabel"] * {
            color: #E5E7EB !important;
        }

        section[data-testid="stSidebar"] div[data-testid="stButton"] > button {
            background: rgba(15, 23, 42, 0.92) !important;
            border-color: rgba(148, 163, 184, 0.28) !important;
            color: #F8FAFC !important;
        }

        section[data-testid="stSidebar"] [data-testid="stExpander"] {
            border: 1px solid var(--hub-border);
            border-radius: 8px;
            background: #FFFFFF !important;
            margin-bottom: 12px;
            box-shadow: 0 4px 14px rgba(15, 23, 42, 0.04);
        }

        [data-testid="stExpander"] summary {
            color: #0F172A !important;
            font-weight: 700 !important;
        }

        [data-testid="stExpander"] summary * {
            color: #0F172A !important;
        }

        [data-testid="stSelectbox"] *,
        [data-testid="stTextInput"] *,
        [data-testid="stTextArea"] *,
        [data-testid="stNumberInput"] *,
        [data-testid="stFileUploader"] *,
        [data-testid="stMultiSelect"] *,
        [data-testid="stRadio"] * {
            color: #0F172A !important;
        }

        [data-testid="stTextInput"],
        [data-testid="stTextArea"],
        [data-testid="stSelectbox"],
        [data-testid="stMultiSelect"] {
            pointer-events: auto !important;
        }

        [data-testid="stTextInput"] *,
        [data-testid="stTextArea"] *,
        [data-testid="stSelectbox"] *,
        [data-testid="stMultiSelect"] * {
            pointer-events: auto !important;
        }

        [data-testid="stSelectbox"] > div,
        [data-testid="stMultiSelect"] > div,
        [data-baseweb="select"],
        [data-baseweb="select"] > div,
        [data-baseweb="select"] div,
        [data-baseweb="popover"],
        [data-baseweb="popover"] div,
        [role="listbox"],
        [role="listbox"] div,
        [role="option"],
        [data-baseweb="menu"],
        [data-baseweb="menu"] div {
            background: #FFFFFF !important;
            color: #0F172A !important;
            border-color: #CBD5E1 !important;
        }

        [data-baseweb="select"] {
            border: 1px solid #CBD5E1 !important;
            border-radius: 12px !important;
            min-height: 44px !important;
        }

        [data-baseweb="select"] input,
        [data-baseweb="select"] span,
        [data-baseweb="select"] svg,
        [data-baseweb="popover"] span,
        [role="listbox"] span,
        [role="option"] span {
            color: #0F172A !important;
            fill: #0F172A !important;
        }

        [data-baseweb="tag"] {
            background: #EFF6FF !important;
            border: 1px solid #BFDBFE !important;
            color: #0F172A !important;
        }

        [data-baseweb="tag"] span,
        [data-baseweb="tag"] svg {
            color: #0F172A !important;
            fill: #0F172A !important;
        }

        .stTextInput > div > div,
        [data-testid="stTextInput"] > div > div,
        [data-baseweb="input"] {
            align-items: center !important;
            background: #FFFFFF !important;
            border: 1px solid #CBD5E1 !important;
            border-radius: 12px !important;
            box-shadow: none !important;
            color: #0F172A !important;
            min-height: 44px !important;
            outline: none !important;
            overflow: hidden !important;
        }

        .stTextInput > div > div:focus-within,
        [data-testid="stTextInput"] > div > div:focus-within,
        [data-baseweb="input"]:focus-within {
            border: 2px solid #2563EB !important;
            box-shadow: none !important;
            outline: none !important;
        }

        [data-baseweb="input"] > div,
        [data-baseweb="input"] button,
        [data-baseweb="input"] [role="button"] {
            align-items: center !important;
            align-self: stretch !important;
            background: #FFFFFF !important;
            border: 0 !important;
            border-radius: 12px !important;
            box-shadow: none !important;
            color: #0F172A !important;
            display: flex !important;
            min-height: 44px !important;
            outline: none !important;
        }

        [data-baseweb="input"] svg {
            color: #0F172A !important;
            fill: #0F172A !important;
        }

        .stTextInput input,
        .stTextInput > div > div > input,
        .stTextArea textarea,
        .stNumberInput input,
        [data-testid="stTextInput"] input,
        [data-baseweb="input"] input,
        [data-testid="stTextArea"] textarea,
        [data-testid="stNumberInput"] input {
            background: #FFFFFF !important;
            border: 0 !important;
            border-radius: 12px !important;
            box-shadow: none !important;
            color: #0F172A !important;
            min-height: 44px;
            outline: none !important;
            padding: 10px 12px !important;
        }

        .stTextArea textarea:hover,
        .stNumberInput input:hover,
        [data-testid="stTextArea"] textarea:hover,
        [data-testid="stNumberInput"] input:hover {
            border: 1px solid #2563EB !important;
        }

        .stTextInput input:focus,
        .stTextInput > div > div > input:focus,
        [data-testid="stTextInput"] input:focus,
        [data-baseweb="input"] input:focus {
            border: 0 !important;
            box-shadow: none !important;
            outline: none !important;
        }

        .stTextArea textarea:focus,
        .stNumberInput input:focus,
        [data-testid="stTextArea"] textarea:focus,
        [data-testid="stNumberInput"] input:focus {
            border: 2px solid #2563EB !important;
            box-shadow: 0 0 0 2px rgba(37, 99, 235, 0.08) !important;
        }

        .stTextInput input::placeholder,
        .stTextArea textarea::placeholder,
        .stNumberInput input::placeholder,
        [data-testid="stTextInput"] input::placeholder,
        [data-testid="stTextArea"] textarea::placeholder,
        [data-testid="stNumberInput"] input::placeholder {
            color: #94A3B8 !important;
            opacity: 1 !important;
        }

        [data-testid="stFileUploader"] {
            background: #F8FAFC !important;
            border: 1px solid #CBD5E1 !important;
            border-radius: 12px !important;
            color: #0F172A !important;
            padding: 12px !important;
        }

        [data-testid="stFileUploader"] section,
        [data-testid="stFileUploader"] section *,
        [data-testid="stFileUploaderDropzone"],
        [data-testid="stFileUploaderDropzone"] * {
            background: #F8FAFC !important;
            color: #0F172A !important;
            border-color: #CBD5E1 !important;
        }

        [data-testid="stFileUploader"] button,
        [data-testid="stFileUploaderDropzone"] button {
            background: #FFFFFF !important;
            border: 1px solid #CBD5E1 !important;
            color: #0F172A !important;
        }

        div[data-testid="stButton"] > button,
        div[data-testid="stDownloadButton"] > button {
            border-radius: 12px;
            border: 1px solid #CBD5E1;
            font-weight: 600;
            width: 100%;
            color: #0F172A !important;
            background: #FFFFFF;
        }

        div[data-testid="stButton"] > button[kind="primary"] {
            background: var(--hub-accent-soft);
            border-color: var(--hub-accent);
            color: #2563EB !important;
            box-shadow: 0 8px 18px rgba(37, 99, 235, 0.24);
        }

        .hub-hero {
            background:
                radial-gradient(circle at top left, rgba(37, 99, 235, 0.10), transparent 34%),
                radial-gradient(circle at top right, rgba(124, 58, 237, 0.10), transparent 30%),
                #FFFFFF;
            border: 1px solid var(--hub-border);
            border-radius: 8px;
            padding: 34px 34px;
            margin: 8px 0 18px;
            box-shadow: 0 14px 40px rgba(15, 23, 42, 0.07);
            max-width: 100%;
            width: 100%;
        }

        .hub-hero h1 {
            color: #0F172A !important;
            font-size: clamp(24px, 3rem, 48px);
            line-height: 1.08;
            margin: 0 0 8px;
            font-weight: 800;
        }

        .hub-hero h2 {
            color: #2563EB !important;
            font-size: clamp(18px, 2rem, 32px);
            margin: 0 0 12px;
            font-weight: 700 !important;
        }

        .hero-subtitle {
            color: #2563EB !important;
            font-size: clamp(18px, 2rem, 32px);
            font-weight: 700 !important;
            line-height: 1.4;
        }

        .hub-hero p {
            color: #475569 !important;
            font-size: clamp(14px, 1rem, 16px);
            line-height: 1.55;
            max-width: 100%;
            margin: 0;
        }

        .status-card-native {
            background: #FFFFFF;
            border: 1px solid var(--hub-border);
            border-radius: 12px;
            box-sizing: border-box;
            min-height: 190px;
            padding: 16px;
            box-shadow: 0 8px 24px rgba(15, 23, 42, 0.04);
        }

        .status-card-top {
            align-items: center;
            display: flex;
            gap: 12px;
            justify-content: space-between;
            margin-bottom: 18px;
        }

        .status-card-icon {
            align-items: center;
            background: #EEF2FF;
            border-radius: 12px;
            display: flex;
            flex-shrink: 0;
            font-size: 22px;
            height: 44px;
            justify-content: center;
            width: 2.75rem;
        }

        .status-card-badge {
            background: #DCFCE7;
            border: 1px solid #86EFAC;
            border-radius: 999px;
            color: #15803D !important;
            flex-shrink: 0;
            font-size: 12px;
            font-weight: 800;
            padding: 6px 12px;
            white-space: nowrap;
        }

        .status-card-title {
            color: #0F172A !important;
            font-size: 18px;
            font-weight: 800;
            line-height: 1.25;
            margin-bottom: 14px;
            hyphens: none;
            overflow-wrap: normal;
            word-break: normal;
        }

        .status-card-description {
            color: #475569 !important;
            font-size: 14px;
            line-height: 1.55;
            hyphens: none;
            overflow-wrap: normal;
            word-break: normal;
        }

        @media (hover: hover) and (pointer: fine) {
            .status-card-native:hover {
                box-shadow: 0 12px 32px rgba(0, 0, 0, 0.08);
                transform: translateY(-4px);
            }
        }

        .status-card-native {
            transition: transform 0.2s ease, box-shadow 0.2s ease;
        }

        .mobile-guide-wrapper {
            display: block;
        }

        .mobile-guide-card {
            background: #FFFFFF;
            border: 1px solid #E2E8F0;
            border-radius: 16px;
            color: #0F172A !important;
            margin-bottom: 12px;
            padding: 16px;
        }

        .mobile-guide-title {
            color: #0F172A !important;
            font-size: 1rem;
            font-weight: 800;
            margin-bottom: 6px;
        }

        .mobile-guide-text {
            color: #475569 !important;
            font-size: 0.92rem;
            line-height: 1.5;
        }

        .mobile-guide-step {
            color: #2563EB !important;
            font-weight: 800;
        }

        .auth-panel {
            background: #FFFFFF;
            border: 1px solid #E2E8F0;
            border-radius: 16px;
            padding: 24px;
        }

        .auth-title {
            color: #0F172A !important;
            font-size: 28px;
            font-weight: 800;
            line-height: 1.2;
            margin-bottom: 6px;
        }

        .auth-subtitle {
            color: #475569 !important;
            font-size: 15px;
            line-height: 1.5;
            margin-bottom: 18px;
        }

        .auth-coming-soon {
            color: #64748B !important;
            font-size: 13px;
            margin-top: 8px;
        }

        .hub-section-title {
            color: #0F172A !important;
            font-size: 19px;
            font-weight: 780;
            margin: 20px 0 8px;
        }

        .hub-result-body {
            background: #FFFFFF;
            border: 1px solid var(--hub-border);
            border-radius: 8px;
            padding: 18px;
            line-height: 1.72;
            color: #0F172A !important;
            white-space: pre-wrap;
            font-size: 15px;
            box-shadow: inset 0 1px 0 rgba(15, 23, 42, 0.02);
        }

        [data-testid="stMetric"] {
            background: #FFFFFF;
            border: 1px solid var(--hub-border);
            border-radius: 8px;
            padding: 14px 16px;
        }

        [data-testid="stMetricLabel"] {
            color: #475569 !important;
            font-weight: 650;
        }

        [data-testid="stMetricValue"] {
            color: #0F172A !important;
            font-weight: 780;
        }

        .subscription-mobile-card {
            display: none;
        }

        .subscription-desktop-metrics {
            display: grid;
            gap: 0.5rem;
        }

        .subscription-desktop-metric {
            background: #FFFFFF;
            border: 1px solid var(--hub-border);
            border-radius: 8px;
            padding: 14px 16px;
        }

        .subscription-desktop-label {
            color: #475569 !important;
            font-size: 14px;
            font-weight: 650;
            line-height: 1.25;
            margin-bottom: 4px;
        }

        .subscription-desktop-value {
            color: #0F172A !important;
            font-size: 24px;
            font-weight: 780;
            line-height: 1.25;
        }

        .hub-mobile-tabs {
            display: none;
            margin: 8px 0 14px;
        }

        .hub-mobile-tabs .tab-row {
            display: flex;
            gap: 8px;
            overflow-x: auto;
            padding-bottom: 4px;
        }

        .hub-mobile-tabs .tab-pill {
            background: #FFFFFF;
            border: 1px solid #E2E8F0;
            border-radius: 999px;
            color: #0F172A !important;
            flex: 0 0 auto;
            font-size: 13px;
            font-weight: 700;
            padding: 8px 11px;
        }

        .workspace-title {
            color: #0F172A !important;
            font-size: 3rem;
            font-weight: 700;
            line-height: 1.08;
            margin: 0 0 0.7rem;
        }

        .workspace-subtitle {
            color: #4f46e5 !important;
            font-size: 1.7rem;
            font-weight: 600;
            line-height: 1.3;
            margin-bottom: 0;
        }

        .hero-title {
            color: #0F172A !important;
            font-size: 3rem;
            font-weight: 700;
            line-height: 1.08;
            margin: 0 0 0.7rem;
        }

        .hero-subtitle {
            font-size: 1.7rem;
            font-weight: 600;
            color: #4f46e5 !important;
            line-height: 1.3;
            margin: 0;
        }

        .hero-container {
            padding: 2rem 0 1rem 0;
        }

        .hero-brand {
            font-size: 3.2rem;
            font-weight: 800;
            line-height: 1.05;
            background: linear-gradient(90deg, #4f46e5, #06b6d4);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            color: transparent !important;
        }

        .hero-sub {
            font-size: 1.2rem;
            color: #6b7280 !important;
            margin-top: 0.5rem;
        }

        .workspace-hero {
            background: #FFFFFF;
            border: 1px solid #E2E8F0;
            border-radius: 8px;
            box-shadow: 0 14px 38px rgba(15, 23, 42, 0.06);
            margin: 0 0 2rem;
            padding: 2.35rem 2.5rem;
            width: 100%;
        }

        .workspace-section-title {
            color: #0F172A !important;
            font-size: clamp(18px, 1.25rem, 24px);
            font-weight: 800;
            margin: 2.2rem 0 1rem;
        }

        .dashboard-metric-card {
            background: linear-gradient(145deg, #ffffff, #f3f4f6);
            border: 1px solid rgba(226, 232, 240, 0.78);
            border-radius: 16px;
            box-shadow: 0 8px 20px rgba(0, 0, 0, 0.05);
            min-height: 132px;
            padding: 20px;
            transition: transform 0.2s ease, box-shadow 0.2s ease;
        }

        .dashboard-metric-card:hover {
            box-shadow: 0 14px 30px rgba(15, 23, 42, 0.08);
            transform: translateY(-4px);
        }

        .dashboard-metric-label {
            color: #64748B !important;
            font-size: 0.88rem;
            font-weight: 700;
            line-height: 1.25;
            margin-bottom: 0.8rem;
        }

        .dashboard-metric-value {
            color: #0F172A !important;
            font-size: 2rem;
            font-weight: 800;
            line-height: 1.1;
        }

        .dashboard-metric-note {
            color: #64748B !important;
            font-size: 0.82rem;
            font-weight: 600;
            margin-top: 0.8rem;
        }

        .quick-action-card {
            background: linear-gradient(135deg, #4f46e5, #6366f1);
            border: 1px solid rgba(255, 255, 255, 0.20);
            border-radius: 16px;
            box-shadow: 0 12px 26px rgba(79, 70, 229, 0.22);
            color: white !important;
            margin-bottom: 10px;
            min-height: 112px;
            padding: 18px;
        }

        .quick-action-title {
            color: white !important;
            font-size: 1.1rem;
            font-weight: 600;
            line-height: 1.25;
            margin-bottom: 0.45rem;
        }

        .quick-action-description {
            color: white !important;
            font-size: 0.9rem;
            line-height: 1.45;
            opacity: 0.85;
        }

        .dashboard-chart-card,
        .copilot-panel {
            background: #FFFFFF;
            border: 1px solid #E2E8F0;
            border-radius: 14px;
            box-shadow: 0 10px 25px rgba(0, 0, 0, 0.06);
            padding: 18px;
            width: 100%;
        }

        .dashboard-chart-title {
            color: #0F172A !important;
            font-size: 1rem;
            font-weight: 800;
            margin-bottom: 0.7rem;
        }

        .copilot-panel {
            position: sticky;
            top: 1rem;
        }

        .copilot-title {
            color: #0F172A !important;
            font-size: 1.25rem;
            font-weight: 800;
            margin-bottom: 0.35rem;
        }

        .copilot-subtitle {
            color: #64748B !important;
            font-size: 0.88rem;
            line-height: 1.45;
            margin-bottom: 1rem;
        }

        .workspace-activity-card,
        .workspace-overview-card {
            background: #FFFFFF;
            border: 1px solid #E2E8F0;
            border-radius: 8px;
            box-shadow: 0 10px 26px rgba(15, 23, 42, 0.05);
            padding: 1.25rem;
            width: 100%;
            max-width: 100%;
        }

        .workspace-activity-item,
        .workspace-overview-line {
            color: #475569 !important;
            font-size: 14px;
            line-height: 1.55;
            margin: 6px 0;
        }

        .workspace-overview-line strong {
            color: #0F172A !important;
        }

        .template-card {
            background: #FFFFFF;
            border: 1px solid #E2E8F0;
            border-radius: 12px;
            padding: 16px;
            width: 100%;
            max-width: 100%;
        }

        .template-card-title {
            color: #0F172A !important;
            font-size: 18px;
            font-weight: 800;
            line-height: 1.25;
            margin-bottom: 6px;
        }

        .template-card-description {
            color: #475569 !important;
            font-size: 14px;
            line-height: 1.5;
            margin-bottom: 10px;
        }

        .template-card-meta {
            color: #475569 !important;
            font-size: 13px;
            font-weight: 650;
            margin-bottom: 10px;
        }

        .template-card-price {
            color: #2563EB !important;
            font-size: 13px;
            font-weight: 800;
        }

        div[data-testid="stVerticalBlockBorderWrapper"] {
            background: #FFFFFF;
            border-color: #E2E8F0 !important;
            box-shadow: 0 8px 24px rgba(15, 23, 42, 0.04);
        }

        @media (max-width: 768px) {
            .block-container {
                padding: 12px !important;
                max-width: 100% !important;
            }

            .hub-hero {
                padding: 16px;
                margin-top: 20px;
                margin-bottom: 20px;
            }

            .hub-mobile-tabs {
                display: block;
                margin-top: 20px;
                margin-bottom: 20px;
            }

            .mobile-guide-wrapper {
                display: block;
                margin-top: 20px;
                margin-bottom: 20px;
            }

            .mobile-guide-card,
            [data-testid="stExpander"],
            div[data-testid="stVerticalBlockBorderWrapper"],
            .hub-result-body {
                padding: 16px !important;
                margin-top: 20px !important;
                margin-bottom: 20px !important;
            }

            .subscription-desktop-metrics {
                display: none !important;
            }

            .subscription-mobile-card {
                background: #FFFFFF;
                border: 1px solid #E2E8F0;
                border-radius: 12px;
                box-sizing: border-box;
                color: #475569 !important;
                display: block;
                margin: 8px 0 14px;
                max-height: 110px;
                min-height: 0;
                overflow: hidden;
                padding: 12px;
            }

            .subscription-mobile-card * {
                color: #475569 !important;
            }

            .subscription-mobile-title {
                color: #0F172A !important;
                font-size: 15px;
                font-weight: 800;
                line-height: 1.15;
                margin-bottom: 6px;
            }

            .subscription-mobile-line {
                color: #475569 !important;
                font-size: 12px;
                font-weight: 650;
                line-height: 1.25;
                margin: 2px 0;
            }

            .subscription-mobile-renewal {
                color: #475569 !important;
                font-size: 12px;
                line-height: 1.25;
                margin-top: 5px;
            }

            .subscription-mobile-renewal strong {
                color: #0F172A !important;
                font-weight: 800;
            }

            .status-card-native {
                min-height: auto;
            }

            .status-card-title {
                font-size: 17px;
            }

            .status-card-description {
                font-size: 14px;
            }

            section[data-testid="stSidebar"] {
                width: min(92vw, 24rem) !important;
            }

            div[data-testid="stHorizontalBlock"] {
                gap: 0.7rem !important;
            }

            div[data-testid="column"] {
                min-width: 0 !important;
                width: 100% !important;
            }

            div[data-testid="stButton"] > button,
            div[data-testid="stDownloadButton"] > button {
                width: 100% !important;
                min-height: 56px !important;
                border-radius: 16px !important;
            }

            textarea {
                min-height: 220px !important;
            }

            .stTextInput > div > div,
            [data-testid="stTextInput"] > div > div,
            [data-baseweb="input"],
            [data-baseweb="input"] > div,
            [data-baseweb="input"] button,
            [data-baseweb="input"] [role="button"] {
                min-height: 48px !important;
            }

            .stTextInput input,
            .stTextArea textarea,
            .stNumberInput input,
            [data-testid="stTextInput"] input,
            [data-testid="stTextArea"] textarea,
            [data-testid="stNumberInput"] input {
                min-height: 48px;
                width: 100%;
            }

            [data-testid="stExpander"] {
                overflow: hidden;
            }

            [data-testid="stSelectbox"],
            [data-testid="stMultiSelect"],
            [data-testid="stTextInput"],
            [data-testid="stTextArea"],
            [data-testid="stFileUploader"] {
                margin-top: 12px !important;
                margin-bottom: 12px !important;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_hero() -> None:
    st.markdown(
        """
        <div class="hub-hero">
            <h1>Gao Intelligence Hub</h1>
            <h2 class="hero-subtitle">Business, Construction &amp; Executive AI Intelligence Platform</h2>
            <p>Transform project documents, business data, and executive workflows into decision-ready intelligence.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_mobile_guide() -> None:
    steps = [
        (
            "1",
            "Open the sidebar",
            "Tap the arrow/menu on the left to access tools.",
        ),
        (
            "2",
            "Choose a tool",
            "Use Prompt Library, Form Builder, Document Intelligence, Automation Intelligence, or Executive Intelligence.",
        ),
        (
            "3",
            "Generate your request",
            "Load a template or generate a professional prompt.",
        ),
        (
            "4",
            "Submit and download",
            "Run the analysis, then download the PDF report.",
        ),
    ]
    with st.expander("📱 Mobile Quick Guide", expanded=False):
        for step_number, title, text in steps:
            st.markdown(
                f"""
                <div class="mobile-guide-wrapper">
                    <div class="mobile-guide-card">
                        <div class="mobile-guide-title">
                            <span class="mobile-guide-step">{step_number}.</span> {title}
                        </div>
                        <div class="mobile-guide-text">{text}</div>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )


def render_status_cards() -> None:
    cards = [
        {
            "icon": "🔬",
            "title": "Business Intelligence",
            "description": "Market, sales, financial, and operating intelligence.",
        },
        {
            "icon": "🏗️",
            "title": "Construction Intelligence",
            "description": "Project, contract, compliance, and delivery analysis.",
        },
        {
            "icon": "📄",
            "title": "Document Intelligence",
            "description": "Extract and analyze PDF, DOCX, XLSX, and TXT files.",
        },
        {
            "icon": "🧠",
            "title": "Executive Decision Support",
            "description": "Board-ready reports powered by executive agent modes.",
        },
    ]
    card_markup = "\n".join(
        f"""
        <div class="status-card-native card">
            <div class="status-card-top">
                <div class="status-card-icon">{card["icon"]}</div>
                <div class="status-card-badge">Active</div>
            </div>
            <div class="status-card-title">{escape(card["title"])}</div>
            <div class="status-card-description">{escape(card["description"])}</div>
        </div>
        """
        for card in cards
    )
    st.markdown(f'<div class="grid status-card-grid">{card_markup}</div>', unsafe_allow_html=True)


def render_mobile_navigation_tabs() -> None:
    st.markdown(
        """
        <div class="hub-mobile-tabs">
            <div class="tab-row">
                <div class="tab-pill">Prompt Library</div>
                <div class="tab-pill">Form Builder</div>
                <div class="tab-pill">Document Analysis</div>
                <div class="tab-pill">Executive Intelligence</div>
                <div class="tab-pill">Automation Intelligence</div>
                <div class="tab-pill">History</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def initialize_auth_state() -> None:
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "current_user" not in st.session_state:
        st.session_state.current_user = None
    if "user_id" not in st.session_state:
        st.session_state.user_id = None
    if "username" not in st.session_state:
        st.session_state.username = None
    if "role" not in st.session_state:
        st.session_state.role = None


def set_authenticated_user(user: dict) -> None:
    st.session_state.authenticated = True
    st.session_state.current_user = user
    st.session_state.user_id = user["id"]
    st.session_state.username = user["username"]
    st.session_state.role = user["role"]


def clear_authenticated_user() -> None:
    st.session_state.authenticated = False
    st.session_state.current_user = None
    st.session_state.user_id = None
    st.session_state.username = None
    st.session_state.role = None
    st.session_state.active_project_id = None
    st.session_state.selected_history = None


def render_auth_page() -> None:
    st.markdown(
        """
        <div class="hub-hero">
            <h1>Gao Intelligence Hub</h1>
            <h2 class="hero-subtitle">Business, Construction &amp; Executive AI Intelligence Platform</h2>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.container(border=True):
        sign_in_tab, create_account_tab = st.tabs(["Sign In", "Create Account"])

        with sign_in_tab:
            with st.form("sign_in_form"):
                identifier = st.text_input("Username or Email", key="auth_sign_in_identifier")
                password = st.text_input("Password", type="password", key="auth_sign_in_password")
                submitted = st.form_submit_button("Sign In", type="primary", use_container_width=True)

            if submitted:
                try:
                    user = user_store.authenticate_user(identifier, password)
                except user_store.UserStoreError:
                    st.error("Could not sign in right now. Please try again later.")
                else:
                    if user is None:
                        st.error("Invalid username/email or password.")
                    else:
                        set_authenticated_user(user)
                        st.rerun()

        with create_account_tab:
            with st.form("create_account_form"):
                username = st.text_input("Username", key="auth_create_username")
                email = st.text_input("Email", key="auth_create_email")
                password = st.text_input("Password", type="password", key="auth_create_password")
                confirm_password = st.text_input(
                    "Confirm Password",
                    type="password",
                    key="auth_create_confirm_password",
                )
                submitted = st.form_submit_button("Create Account", type="primary", use_container_width=True)

            if submitted:
                if password != confirm_password:
                    st.error("Passwords do not match.")
                else:
                    try:
                        user = user_store.create_user(username, email, password)
                    except ValueError as exc:
                        st.error(str(exc))
                    except user_store.UserStoreError:
                        st.error("Could not create account right now. Please try again later.")
                    else:
                        set_authenticated_user(user)
                        st.rerun()


def get_current_user_id() -> int:
    if st.session_state.user_id is not None:
        return int(st.session_state.user_id)
    current_user = st.session_state.current_user or {}
    return int(current_user["id"])


def format_limit(limit: int | None) -> str:
    if limit is None:
        return "Unlimited"
    return str(limit)


def format_renewal_date(value: str | None) -> str:
    parsed_value = subscription_store.parse_datetime(value)
    if parsed_value is None:
        return "Not set"
    return parsed_value.date().isoformat()


def refresh_subscription_usage() -> dict:
    usage = subscription_store.get_usage(get_current_user_id())
    if st.session_state.current_user is not None:
        st.session_state.current_user.update(usage)
    return usage


def consume_ai_request_or_warn() -> bool:
    can_use, message = subscription_store.can_use_ai_request(get_current_user_id())
    if not can_use:
        st.error(message)
        return False
    try:
        subscription_store.increment_ai_request(get_current_user_id())
    except PermissionError as exc:
        st.error(str(exc))
        return False
    refresh_subscription_usage()
    return True


def consume_document_analysis_or_warn() -> bool:
    can_use, message = subscription_store.can_use_document_analysis(get_current_user_id())
    if not can_use:
        st.error(message)
        return False
    try:
        subscription_store.increment_document_analysis(get_current_user_id())
    except PermissionError as exc:
        st.error(str(exc))
        return False
    refresh_subscription_usage()
    return True


def consume_executive_agent_usage_or_warn() -> bool:
    can_use_documents, document_message = subscription_store.can_use_document_analysis(get_current_user_id())
    if not can_use_documents:
        st.error(document_message)
        return False

    can_use_ai, ai_message = subscription_store.can_use_ai_request(get_current_user_id())
    if not can_use_ai:
        st.error(ai_message)
        return False

    try:
        subscription_store.increment_document_analysis(get_current_user_id())
        subscription_store.increment_ai_request(get_current_user_id())
    except PermissionError as exc:
        st.error(str(exc))
        return False

    refresh_subscription_usage()
    return True


def render_subscription_summary() -> None:
    try:
        usage = refresh_subscription_usage()
    except ValueError:
        st.error("Could not load subscription details.")
        return

    request_usage = f'{usage["current_request_count"]} / {format_limit(usage["monthly_request_limit"])}'
    document_usage = f'{usage["current_document_count"]} / {format_limit(usage["monthly_document_limit"])}'
    renewal_date = format_renewal_date(usage["subscription_end"])

    st.markdown("#### Subscription")
    st.markdown(
        f"""
        <div class="subscription-desktop-metrics">
            <div class="subscription-desktop-metric">
                <div class="subscription-desktop-label">Plan</div>
                <div class="subscription-desktop-value">{escape(str(usage["subscription_tier"]))}</div>
            </div>
            <div class="subscription-desktop-metric">
                <div class="subscription-desktop-label">Usage</div>
                <div class="subscription-desktop-value">{escape(request_usage)}</div>
            </div>
            <div class="subscription-desktop-metric">
                <div class="subscription-desktop-label">Documents</div>
                <div class="subscription-desktop-value">{escape(document_usage)}</div>
            </div>
            <div class="subscription-desktop-metric">
                <div class="subscription-desktop-label">Renewal Date</div>
                <div class="subscription-desktop-value">{escape(renewal_date)}</div>
            </div>
        </div>
        <div class="subscription-mobile-card">
            <div class="subscription-mobile-title">{escape(str(usage["subscription_tier"]))}</div>
            <div class="subscription-mobile-line">{escape(request_usage)} Requests</div>
            <div class="subscription-mobile-line">{escape(document_usage)} Documents</div>
            <div class="subscription-mobile-renewal">Renews:<br><strong>{escape(renewal_date)}</strong></div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_user_limit_summary() -> None:
    try:
        active_users, max_users = user_store.get_user_limit_status()
    except user_store.UserStoreError:
        st.error("Could not load user count.")
        return

    max_users_label = "Unlimited" if max_users == 0 else str(max_users)
    st.caption(f"Users: {active_users} / {max_users_label}")


def set_active_section(section: str, documents_view: str | None = None) -> None:
    st.session_state.active_section = section
    section_to_navigation = {
        "workspace": "Dashboard",
        "documents": "Documents",
        "ai": "Agents",
        "automations": "Automations",
        "logs": "Logs",
        "settings": "Settings",
    }
    if section in section_to_navigation:
        st.session_state.nav_selected = section_to_navigation[section]
        st.session_state.last_nav_selected = section_to_navigation[section]
    if documents_view is not None:
        st.session_state.documents_view = documents_view


def get_active_project() -> dict | None:
    active_project_id = st.session_state.get("active_project_id")
    if active_project_id is None:
        return None
    try:
        return project_store.get_project(int(active_project_id), get_current_user_id())
    except (ValueError, project_store.ProjectStoreError):
        return None


def require_active_project() -> bool:
    if get_active_project() is None:
        st.warning("Please create or select a project first.")
        return False
    return True


def add_active_project_history(action_type: str, content: str) -> None:
    active_project = get_active_project()
    if active_project is None:
        return
    try:
        project_store.add_project_history(active_project["id"], action_type, content)
    except (ValueError, project_store.ProjectStoreError):
        st.warning("Project history could not be saved.")


def render_new_project_form() -> None:
    with st.expander("New Project", expanded=False):
        with st.form("new_project_form"):
            project_name = st.text_input("Project Name", key="new_project_name")
            project_description = st.text_area("Description", key="new_project_description", height=90)
            submitted = st.form_submit_button("Create Project", type="primary", use_container_width=True)

        if submitted:
            try:
                project = project_store.create_project(
                    get_current_user_id(),
                    project_name,
                    project_description,
                )
            except ValueError as exc:
                st.error(str(exc))
            except project_store.ProjectStoreError:
                st.error("Could not create project right now.")
            else:
                st.session_state.active_project_id = project["id"]
                st.session_state.active_section = "projects"
                st.rerun()


def render_projects_list(title: str = "Recent Projects") -> None:
    st.markdown(f'<div class="workspace-section-title">{escape(title)}</div>', unsafe_allow_html=True)
    try:
        projects = project_store.list_projects(get_current_user_id())
    except project_store.ProjectStoreError:
        st.error("Could not load projects.")
        return

    if not projects:
        st.caption("No projects yet.")
        return

    for project in projects:
        project_label = project["name"]
        if project.get("description"):
            project_label = f'{project_label} - {project["description"]}'
        if st.button(project_label, key=f'project_select_{project["id"]}', use_container_width=True):
            st.session_state.active_project_id = project["id"]
            st.session_state.active_section = "projects"
            st.rerun()


def render_project_history(project_id: int) -> None:
    st.markdown('<div class="workspace-section-title">History Timeline</div>', unsafe_allow_html=True)
    try:
        history_items = project_store.list_project_history(project_id)
    except project_store.ProjectStoreError:
        st.error("Could not load project history.")
        return

    if not history_items:
        st.caption("No project history yet.")
        return

    for item in history_items:
        st.markdown(
            f'<div class="workspace-activity-item">{escape(item["content"])}</div>',
            unsafe_allow_html=True,
        )


def render_project_view(project: dict) -> None:
    st.markdown(
        f"""
        <div class="workspace-title">{escape(project["name"])}</div>
        <div class="workspace-subtitle">{escape(project.get("description") or "No description")}</div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown('<div class="workspace-section-title">Quick Actions</div>', unsafe_allow_html=True)
    if st.button("Analyze Document", key="project_analyze_document", use_container_width=True):
        set_active_section("documents", "document")
        st.rerun()
    if st.button("Build Prompt", key="project_build_prompt", use_container_width=True):
        set_active_section("ai")
        st.rerun()
    if st.button("Create Automation", key="project_create_automation", use_container_width=True):
        set_active_section("automations")
        st.rerun()
    if st.button("Executive Analysis", key="project_executive_analysis", use_container_width=True):
        set_active_section("documents", "executive")
        st.rerun()

    render_project_history(project["id"])


def render_projects_page() -> None:
    active_project = get_active_project()
    if active_project is not None:
        render_project_view(active_project)
        st.markdown('<div class="workspace-section-title">Projects</div>', unsafe_allow_html=True)

    render_new_project_form()
    render_projects_list("Projects List")


def render_workspace() -> None:
    current_user = st.session_state.current_user or {}
    username = st.session_state.username or current_user.get("username", "User")
    try:
        usage = refresh_subscription_usage()
        documents_processed = str(usage["current_document_count"])
        api_requests = str(usage["current_request_count"])
    except (KeyError, ValueError):
        documents_processed = "0"
        api_requests = "0"

    st.markdown(
        f"""
        <div class="hero-container">
            <div class="hero-brand">Gao Intelligence Hub</div>
            <div class="hero-sub">Welcome back, {escape(str(username))}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    metric_columns = st.columns(4)
    metrics = [
        ("Active Agents", str(len(EXECUTIVE_AGENT_MODES)), "Executive review modes ready"),
        ("Documents Processed", documents_processed, "This billing period"),
        ("API Requests", api_requests, "This billing period"),
        ("System Status", "Online", "Core services available"),
    ]
    for column, (label, value, note) in zip(metric_columns, metrics):
        with column:
            st.markdown(
                f"""
                <div class="dashboard-metric-card">
                    <div class="dashboard-metric-label">{escape(label)}</div>
                    <div class="dashboard-metric-value">{escape(value)}</div>
                    <div class="dashboard-metric-note">{escape(note)}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    chart_columns = st.columns(2)
    with chart_columns[0]:
        st.markdown(
            """
            <div class="dashboard-chart-card">
                <div class="dashboard-chart-title">Usage Trend</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.line_chart(
            {
                "Documents": [2, 5, 8, 12, 20],
                "API Requests": [10, 30, 45, 60, 90],
            }
        )

    with chart_columns[1]:
        st.markdown(
            """
            <div class="dashboard-chart-card">
                <div class="dashboard-chart-title">Agent Usage</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        st.bar_chart(
            {
                "Legal": [12],
                "Finance": [8],
                "Risk": [5],
            }
        )

    st.markdown('<div class="workspace-section-title">Quick Actions</div>', unsafe_allow_html=True)
    action_columns = st.columns(4)
    actions = [
        (
            "Analyze Document",
            "Review contracts, tenders, reports, and business documents.",
            "quick_analyze_document",
            "documents",
            "document",
        ),
        (
            "Upload File",
            "Open the document workspace and upload source files.",
            "quick_upload_file",
            "documents",
            "document",
        ),
        (
            "Run Agent",
            "Start an executive agent review for a selected project.",
            "quick_run_agent",
            "documents",
            "executive",
        ),
        (
            "View Logs",
            "Review saved activity and generated intelligence history.",
            "quick_view_logs",
            "settings",
            None,
        ),
    ]

    for column, (title, description, key, section, documents_view) in zip(action_columns, actions):
        with column:
            st.markdown(
                f"""
                <div class="quick-action-card">
                    <div class="quick-action-title">{escape(title)}</div>
                    <div class="quick-action-description">{escape(description)}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            if st.button(title, key=key, use_container_width=True):
                if section != "settings" and not require_active_project():
                    return
                set_active_section(section, documents_view)
                st.rerun()

    render_new_project_form()
    render_projects_list()

    st.markdown('<div class="workspace-section-title">Recent Activity</div>', unsafe_allow_html=True)
    st.markdown(
        """
        <div class="workspace-activity-card">
            <div class="workspace-activity-item">Document Analysis: Construction Proposal</div>
            <div class="workspace-activity-item">Prompt Generated: Marketing Strategy</div>
            <div class="workspace-activity-item">Automation Created: Project Workflow</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    try:
        usage = refresh_subscription_usage()
    except ValueError:
        usage = {
            "subscription_tier": "Starter",
            "current_request_count": 0,
            "monthly_request_limit": 100,
            "current_document_count": 0,
            "monthly_document_limit": 20,
        }

    st.markdown('<div class="workspace-section-title">Subscription Overview</div>', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="workspace-overview-card">
            <div class="workspace-overview-line"><strong>Plan:</strong> {escape(str(usage["subscription_tier"]))}</div>
            <div class="workspace-overview-line"><strong>Usage:</strong> {escape(str(usage["current_request_count"]))} / {escape(format_limit(usage["monthly_request_limit"]))}</div>
            <div class="workspace-overview-line"><strong>Documents:</strong> {escape(str(usage["current_document_count"]))} / {escape(format_limit(usage["monthly_document_limit"]))}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_prompt_library() -> None:
    categories = prompt_store.list_categories()
    if not categories:
        st.caption("No prompts available.")
        return

    selected_category = st.selectbox("Category", categories, key="prompt_library_category")
    prompts = prompt_store.list_prompts(selected_category)
    prompt_options = {prompt["name"]: prompt["id"] for prompt in prompts}

    selected_prompt_name = st.selectbox(
        "Prompt",
        list(prompt_options.keys()),
        key="prompt_library_prompt",
    )
    selected_prompt_id = prompt_options[selected_prompt_name]

    if st.button("Load Prompt Template", use_container_width=True):
        if not require_active_project():
            return
        selected_prompt = prompt_store.get_prompt(selected_prompt_id)
        if selected_prompt:
            st.session_state.query = selected_prompt["content"]
            add_active_project_history("prompt", f"Prompt generated: {selected_prompt['name']}")


def render_ai_request_panel() -> None:
    st.markdown('<div class="hub-section-title">AI Workspace</div>', unsafe_allow_html=True)
    active_project = get_active_project()
    if active_project is None:
        st.warning("Please create or select a project first.")
        return
    st.caption(f'Project: {active_project["name"]}')

    with st.container(border=True):
        mode = st.radio(
            "Routing Mode",
            ["Fast Mode", "Advanced Agent Routing"],
            index=0,
            horizontal=True,
        )
        is_fast_mode = mode == "Fast Mode"
        mode_name = "Fast" if is_fast_mode else "Advanced"
        st.caption(f"Active mode: {mode_name}")

        query = st.text_area(
            "Request",
            height=220,
            placeholder="Enter a request for the gateway...",
            key="query",
        )
        submitted = st.button("Submit", type="primary", use_container_width=True)

    display_entry = st.session_state.selected_history

    if submitted:
        cleaned_query = query.strip()

        if not cleaned_query:
            st.warning("Please enter a request.")
        elif not require_active_project():
            pass
        elif not consume_ai_request_or_warn():
            pass
        else:
            try:
                with st.spinner("Analyzing..."):
                    started_at = time.perf_counter()
                    if is_fast_mode:
                        payload = run_fast_mode(cleaned_query)
                    else:
                        payload = run_advanced_mode(cleaned_query)
                    elapsed_seconds = time.perf_counter() - started_at
            except requests.exceptions.RequestException as exc:
                st.error(f"Request failed: {exc}")
            except ValueError:
                st.error("Request failed: FastAPI returned invalid JSON.")
            else:
                display_entry = save_history_entry(
                    cleaned_query,
                    mode_name,
                    payload,
                    elapsed_seconds,
                )
                add_active_project_history("ai", f"Prompt generated: {cleaned_query[:80]}")

    if display_entry:
        display_result(display_entry)


def render_documents_page() -> None:
    st.markdown('<div class="workspace-title">Documents</div>', unsafe_allow_html=True)
    active_project = get_active_project()
    if active_project is None:
        st.warning("Please create or select a project first.")
        return
    st.caption(f'Project: {active_project["name"]}')

    view_options = ["document", "executive"]
    selected_view = st.radio(
        "Documents View",
        view_options,
        index=view_options.index(st.session_state.get("documents_view", "document")),
        format_func=lambda value: "Document Intelligence" if value == "document" else "Executive Analysis",
        horizontal=True,
    )
    st.session_state.documents_view = selected_view

    if selected_view == "executive":
        render_project_intelligence_review()
    else:
        render_document_analysis()


def render_automation_page() -> None:
    st.markdown('<div class="workspace-title">Automations</div>', unsafe_allow_html=True)
    active_project = get_active_project()
    if active_project is None:
        st.warning("Please create or select a project first.")
        return
    st.caption(f'Project: {active_project["name"]}')
    render_automation_intelligence()


def render_marketplace() -> None:
    st.markdown('<div class="workspace-title">AI Marketplace</div>', unsafe_allow_html=True)
    st.markdown('<div class="workspace-subtitle">Popular Templates</div>', unsafe_allow_html=True)

    search_query = st.text_input("Search", key="marketplace_search")
    selected_category = st.selectbox(
        "Categories",
        ["All", *template_store.TEMPLATE_CATEGORIES],
        key="marketplace_category",
    )

    try:
        templates = template_store.list_templates(selected_category, search_query)
    except template_store.TemplateStoreError:
        st.error("Could not load templates.")
        return

    if not templates:
        st.caption("No templates found.")
        return

    for template in templates:
        price = float(template["price"] or 0)
        price_label = "Free" if price == 0 else f"${price:g}"
        st.markdown(
            f"""
            <div class="template-card">
                <div class="template-card-title">{escape(template["name"])}</div>
                <div class="template-card-description">{escape(template["description"])}</div>
                <div class="template-card-meta">{escape(template["category"])} · <span class="template-card-price">{escape(price_label)}</span></div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        if st.button("Use Template", key=f'use_template_{template["id"]}', use_container_width=True):
            if price > 0:
                st.info("Premium Template (Coming Soon)")
            elif not require_active_project():
                return
            else:
                template_config = template["config_json"]
                st.session_state.query = template_config
                add_active_project_history(
                    "automation",
                    f"Automation created: {template['name']}",
                )
                st.success("Template copied into the current project automation workspace.")


def render_prompt_builder() -> None:
    st.markdown('<div class="workspace-title">AI Center</div>', unsafe_allow_html=True)
    with st.expander("Prompt Library", expanded=True):
        render_prompt_library()
    with st.expander("Form Builder", expanded=False):
        render_form_builder()
    render_ai_request_panel()


def render_settings() -> None:
    st.markdown('<div class="workspace-title">Settings</div>', unsafe_allow_html=True)
    render_subscription_summary()


def render_logs() -> None:
    st.markdown('<div class="workspace-title">Logs</div>', unsafe_allow_html=True)
    st.markdown('<div class="hub-section-title">Intelligence History</div>', unsafe_allow_html=True)
    st.download_button(
        "Export History",
        data=json.dumps(st.session_state.history, indent=2),
        file_name="history.json",
        mime="application/json",
        use_container_width=True,
    )

    if st.button("Clear History", use_container_width=True):
        st.session_state.history = []
        st.session_state.selected_history = None

    if st.session_state.history:
        for index, entry in enumerate(st.session_state.history[:HISTORY_LIMIT]):
            label = entry["query"].replace("\n", " ").strip()
            if len(label) > 60:
                label = f"{label[:57]}..."
            if st.button(
                f"{index + 1}. {entry['mode']} - {label}",
                key=f"settings_history_{index}",
                use_container_width=True,
            ):
                st.session_state.selected_history = entry
                set_active_section("ai")
                st.rerun()
    else:
        st.caption("No history yet.")


def render_ai_copilot_panel() -> None:
    st.markdown(
        """
        <div class="copilot-panel">
            <div class="copilot-title">AI Copilot</div>
            <div class="copilot-subtitle">Ask for summaries, next steps, or workspace guidance.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    user_input = st.text_area("Ask anything", key="copilot_input", height=180)
    if st.button("Run AI", key="copilot_run_ai", use_container_width=True):
        if user_input.strip():
            st.write("AI response placeholder")
        else:
            st.warning("Enter a question first.")


def display_result(entry: dict) -> None:
    st.markdown('<div class="hub-section-title">Response</div>', unsafe_allow_html=True)
    with st.container(border=True):
        result_metrics = [
            ("Mode", str(entry.get("mode", ""))),
            ("Response Time", f"{entry.get('elapsed_seconds', 0.0):.1f}s"),
            ("Category", str(entry.get("category", ""))),
            ("Confidence", str(entry.get("confidence", ""))),
            ("Version", str(entry.get("version", ""))),
        ]
        if entry.get("agent_mode"):
            result_metrics.append(("Agent Mode", str(entry.get("agent_mode", ""))))

        result_metric_markup = "\n".join(
            f"""
            <div class="subscription-desktop-metric">
                <div class="subscription-desktop-label">{escape(label)}</div>
                <div class="subscription-desktop-value">{escape(value)}</div>
            </div>
            """
            for label, value in result_metrics
        )
        st.markdown(f'<div class="grid">{result_metric_markup}</div>', unsafe_allow_html=True)

        st.markdown("#### Result")
        st.markdown(
            f'<div class="hub-result-body">{escape(str(entry.get("result", ""))).replace(chr(10), "<br>")}</div>',
            unsafe_allow_html=True,
        )
    render_pdf_download(entry)


def format_form_value(value: str | list[str]) -> str:
    if isinstance(value, list):
        return ", ".join(value) if value else "Not specified"
    return value.strip() if value.strip() else "Not specified"


def build_prompt_from_form(form_definition: dict, values: dict[str, str | list[str]]) -> str:
    formatted_values = {
        key: format_form_value(value)
        for key, value in values.items()
    }
    return form_definition["template"].format(**formatted_values)


def build_automation_prompt(automation_definition: dict, values: dict[str, str | list[str]]) -> str:
    formatted_values = {
        key: format_form_value(value)
        for key, value in values.items()
    }
    formatted_values["output_sections"] = "\n".join(
        f"{index}. {section}"
        for index, section in enumerate(automation_definition["output_sections"], start=1)
    )
    return automation_definition["template"].format(**formatted_values)


def extract_pdf_text(file_bytes: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(file_bytes))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def extract_xlsx_text(file_bytes: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for worksheet in workbook.worksheets:
        parts.append(f"Sheet: {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            row_values = [str(value) for value in row if value is not None]
            if row_values:
                parts.append(" | ".join(row_values))
    workbook.close()
    return "\n".join(parts)


def extract_uploaded_document_text(file_name: str, file_bytes: bytes) -> str:
    suffix = file_name.rsplit(".", 1)[-1].lower()
    if suffix == "pdf":
        return extract_pdf_text(file_bytes)
    if suffix == "txt":
        return file_bytes.decode("utf-8")
    if suffix == "docx":
        return extract_docx_text(file_bytes)
    if suffix == "xlsx":
        return extract_xlsx_text(file_bytes)
    raise ValueError("Unsupported file type.")


def truncate_document_text(
    document_text: str,
    limit: int = DOCUMENT_TEXT_LIMIT,
) -> tuple[str, bool]:
    if len(document_text) <= limit:
        return document_text, False
    return document_text[:limit], True


def apply_total_document_limit(documents: list[dict]) -> tuple[list[dict], bool]:
    remaining_characters = DOCUMENT_TOTAL_TEXT_LIMIT
    limited_documents = []
    total_truncated = False

    for document in documents:
        document_copy = document.copy()
        text = document_copy["text"]

        if remaining_characters <= 0:
            document_copy["text"] = ""
            document_copy["total_truncated"] = True
            total_truncated = True
        elif len(text) > remaining_characters:
            document_copy["text"] = text[:remaining_characters]
            document_copy["total_truncated"] = True
            total_truncated = True
            remaining_characters = 0
        else:
            document_copy["total_truncated"] = False
            remaining_characters -= len(text)

        limited_documents.append(document_copy)

    return limited_documents, total_truncated


def build_document_prompt(
    analysis_type: str,
    file_name: str,
    document_text: str,
    truncated: bool,
) -> str:
    truncation_notice = ""
    if truncated:
        truncation_notice = f"Content truncated to first {DOCUMENT_TEXT_LIMIT} characters.\n\n"
    report_items = DOCUMENT_REPORT_STRUCTURES.get(
        analysis_type,
        DOCUMENT_REPORT_STRUCTURES["Business Analysis"],
    )
    report_structure = "\n".join(
        f"{index}. {item}"
        for index, item in enumerate(report_items, start=1)
    )

    return f"""Analyze the following uploaded document.

Analysis Type: {analysis_type}
File Name: {file_name}

Document Content:
{truncation_notice}{document_text}

Provide:
{report_structure}"""


def build_project_review_prompt(
    review_type: str,
    documents: list[dict],
    total_truncated: bool,
) -> str:
    documents_text = format_project_documents_for_prompt(documents)
    total_notice = ""
    if total_truncated:
        total_notice = f"\nContent truncated to first {DOCUMENT_TOTAL_TEXT_LIMIT} total characters across uploaded files.\n"

    return f"""Analyze the following project documents as a professional construction and business advisor.

Review Type: {review_type}

Project Documents:
{documents_text}{total_notice}

Provide:
1. Executive Summary
2. Document-by-Document Findings
3. Cross-Document Inconsistencies
4. Commercial Risks
5. Technical Risks
6. Contractual Risks
7. Compliance / NCC Risks
8. Procurement Risks
9. Priority Matrix
10. Recommended Actions"""


def format_project_documents_for_prompt(documents: list[dict]) -> str:
    document_sections = []
    for document in documents:
        notices = []
        if document.get("file_truncated"):
            notices.append(f"File content truncated to first {DOCUMENT_TEXT_LIMIT} characters.")
        if document.get("total_truncated"):
            notices.append(f"Combined document content truncated to first {DOCUMENT_TOTAL_TEXT_LIMIT} characters.")

        notice_text = ""
        if notices:
            notice_text = "\n".join(notices) + "\n"

        document_sections.append(
            f"=== File: {document['file_name']} ===\n{notice_text}{document['text']}"
        )
    return "\n\n".join(document_sections)


def render_document_analysis() -> None:
    uploaded_file = st.file_uploader(
        "Upload Document",
        type=["pdf", "txt", "docx", "xlsx"],
        key="document_analysis_upload",
    )
    analysis_type = st.selectbox(
        "Analysis Type",
        DOCUMENT_ANALYSIS_TYPES,
        key="document_analysis_type",
    )

    document_text = ""
    truncated_text = ""
    truncated = False
    read_failed = False

    if uploaded_file:
        try:
            document_text = extract_uploaded_document_text(
                uploaded_file.name,
                uploaded_file.getvalue(),
            )
            truncated_text, truncated = truncate_document_text(document_text)
        except Exception as exc:
            read_failed = True
            st.error(f"Could not read uploaded file: {exc}")
        else:
            st.write(f"File name: {uploaded_file.name}")
            st.write(f"Extracted character count: {len(document_text)}")
            st.write(f"Truncated: {'Yes' if truncated else 'No'}")
            if truncated:
                st.warning(f"File content was truncated to first {DOCUMENT_TEXT_LIMIT} characters.")

    if st.button("Generate Document Intelligence Prompt", use_container_width=True):
        if not uploaded_file:
            st.warning("Please upload a document first.")
        elif read_failed:
            st.error("Could not generate a prompt because the uploaded file could not be read.")
        elif not require_active_project():
            return
        elif not consume_document_analysis_or_warn():
            return
        else:
            st.session_state.query = build_document_prompt(
                analysis_type,
                uploaded_file.name,
                truncated_text,
                truncated,
            )
            add_active_project_history("document", f"Document analyzed: {uploaded_file.name}")


def render_project_intelligence_review() -> None:
    uploaded_files = st.file_uploader(
        "Upload Project Documents",
        type=["pdf", "txt", "docx", "xlsx"],
        key="project_review_upload",
        accept_multiple_files=True,
    )
    review_type = st.selectbox(
        "Review Type",
        PROJECT_REVIEW_TYPES,
        key="project_review_type",
    )
    agent_mode = st.selectbox(
        "Agent Mode",
        EXECUTIVE_AGENT_MODES,
        key="executive_agent_mode",
    )
    st.caption(EXECUTIVE_AGENT_MODE_DESCRIPTIONS[agent_mode])

    documents = []
    had_file_error = False

    if uploaded_files:
        for uploaded_file in uploaded_files:
            try:
                document_text = extract_uploaded_document_text(
                    uploaded_file.name,
                    uploaded_file.getvalue(),
                )
                truncated_text, file_truncated = truncate_document_text(document_text)
            except Exception as exc:
                had_file_error = True
                st.error(f"Could not read {uploaded_file.name}: {exc}")
            else:
                st.write(f"File name: {uploaded_file.name}")
                st.write(f"Extracted character count: {len(document_text)}")
                st.write(f"Truncated: {'Yes' if file_truncated else 'No'}")
                documents.append(
                    {
                        "file_name": uploaded_file.name,
                        "text": truncated_text,
                        "file_truncated": file_truncated,
                    }
                )

        if any(document["file_truncated"] for document in documents):
            st.warning(f"One or more files were truncated to first {DOCUMENT_TEXT_LIMIT} characters.")

    ready_documents = []
    ready_total_truncated = False
    ready_documents_text = ""
    if documents:
        ready_documents, ready_total_truncated = apply_total_document_limit(documents)
        ready_documents_text = format_project_documents_for_prompt(ready_documents)

    if st.button("Generate Executive Review", use_container_width=True):
        if not uploaded_files:
            st.warning("Please upload at least one project document first.")
        elif not documents:
            st.error("No uploaded project documents could be read.")
        elif not require_active_project():
            return
        elif not consume_document_analysis_or_warn():
            return
        else:
            if ready_total_truncated:
                st.warning(f"Combined document content was truncated to first {DOCUMENT_TOTAL_TEXT_LIMIT} characters.")
            if had_file_error:
                st.warning("Some files could not be read and were excluded from the prompt.")
            st.session_state.query = build_project_review_prompt(
                review_type,
                ready_documents,
                ready_total_truncated,
            )
            add_active_project_history("document", f"Document analyzed: {review_type}")

    if st.button("Run Executive Agent Team", use_container_width=True):
        if not uploaded_files:
            st.warning("Please upload at least one project document first.")
        elif not documents:
            st.error("No uploaded project documents could be read.")
        elif not require_active_project():
            return
        elif not consume_executive_agent_usage_or_warn():
            return
        else:
            if ready_total_truncated:
                st.warning(f"Combined document content was truncated to first {DOCUMENT_TOTAL_TEXT_LIMIT} characters.")
            if had_file_error:
                st.warning("Some files could not be read and were excluded from the agent team input.")
            try:
                from executive_agents import run_executive_agent_team

                with st.spinner("Running executive agent team..."):
                    started_at = time.perf_counter()
                    payload = run_executive_agent_team(
                        review_type,
                        ready_documents_text,
                        agent_mode,
                    )
                    elapsed_seconds = time.perf_counter() - started_at
            except Exception:
                payload = {
                    "category": "executive",
                    "confidence": 0.0,
                    "result": "Executive agent team failed: unable to reach the LLM provider. Please try again later.",
                    "version": VERSION,
                    "agent_mode": agent_mode,
                }
                elapsed_seconds = 0.0

            st.session_state.query = build_project_review_prompt(
                review_type,
                ready_documents,
                ready_total_truncated,
            )
            save_history_entry(
                st.session_state.query,
                "Executive Agent Team",
                payload,
                elapsed_seconds,
                {"agent_mode": agent_mode},
            )
            add_active_project_history("executive", f"Executive analysis: {review_type}")


def render_form_builder() -> None:
    selected_category = st.selectbox(
        "Form Category",
        list(FORM_LIBRARY.keys()),
        key="form_builder_category",
    )
    form_definitions = FORM_LIBRARY[selected_category]
    form_names = [form_definition["name"] for form_definition in form_definitions]
    selected_form_name = st.selectbox(
        "Form",
        form_names,
        key="form_builder_form",
    )
    selected_form = next(
        form_definition
        for form_definition in form_definitions
        if form_definition["name"] == selected_form_name
    )

    with st.form("prompt_form_builder"):
        values = {}
        for form_field in selected_form["fields"]:
            widget_key = f"form_builder_{selected_category}_{selected_form_name}_{form_field['key']}"
            field_type = form_field["type"]

            if field_type == "text_area":
                values[form_field["key"]] = st.text_area(
                    form_field["name"],
                    key=widget_key,
                    height=80,
                    placeholder=form_field.get("placeholder", ""),
                )
            elif field_type == "selectbox":
                values[form_field["key"]] = st.selectbox(
                    form_field["name"],
                    form_field["options"],
                    key=widget_key,
                )
            elif field_type == "multiselect":
                values[form_field["key"]] = st.multiselect(
                    form_field["name"],
                    form_field["options"],
                    key=widget_key,
                )
            else:
                values[form_field["key"]] = st.text_input(
                    form_field["name"],
                    key=widget_key,
                    placeholder=form_field.get("placeholder", ""),
                )

        if st.form_submit_button("Generate Professional Prompt", use_container_width=True):
            if not require_active_project():
                return
            st.session_state.query = build_prompt_from_form(selected_form, values)
            add_active_project_history("prompt", f"Prompt generated: {selected_form_name}")


def render_automation_intelligence() -> None:
    automation_names = [automation_definition["name"] for automation_definition in AUTOMATION_LIBRARY]
    selected_automation_name = st.selectbox(
        "Automation Type",
        automation_names,
        key="automation_intelligence_type",
    )
    selected_automation = next(
        automation_definition
        for automation_definition in AUTOMATION_LIBRARY
        if automation_definition["name"] == selected_automation_name
    )
    selected_automation_key = "".join(
        character.lower() if character.isalnum() else "_"
        for character in selected_automation_name
    ).strip("_")

    with st.form("automation_intelligence_form"):
        values = {}
        for automation_field in selected_automation["fields"]:
            widget_key = f"automation_{selected_automation_key}_{automation_field['key']}"
            field_type = automation_field["type"]

            if field_type == "text_area":
                values[automation_field["key"]] = st.text_area(
                    automation_field["name"],
                    key=widget_key,
                    height=90,
                    placeholder=automation_field.get("placeholder", ""),
                )
            elif field_type == "selectbox":
                values[automation_field["key"]] = st.selectbox(
                    automation_field["name"],
                    automation_field["options"],
                    key=widget_key,
                )
            elif field_type == "multiselect":
                values[automation_field["key"]] = st.multiselect(
                    automation_field["name"],
                    automation_field["options"],
                    key=widget_key,
                )
            else:
                values[automation_field["key"]] = st.text_input(
                    automation_field["name"],
                    key=widget_key,
                    placeholder=automation_field.get("placeholder", ""),
                )

        if st.form_submit_button("Generate Automation Blueprint", use_container_width=True):
            if not require_active_project():
                return
            st.session_state.query = build_automation_prompt(selected_automation, values)
            add_active_project_history("automation", f"Automation created: {selected_automation_name}")


st.set_page_config(
    page_title="Gao Intelligence Hub",
    layout="wide",
    initial_sidebar_state="expanded",
)
inject_custom_css()
initialize_auth_state()
try:
    user_store.initialize_user_store()
except user_store.UserStoreError:
    render_auth_page()
    st.error("Could not start the user database. Please contact the administrator.")
    st.stop()

try:
    project_store.initialize_project_store()
except project_store.ProjectStoreError:
    render_auth_page()
    st.error("Could not start the project database. Please contact the administrator.")
    st.stop()

try:
    template_store.initialize_template_store()
except template_store.TemplateStoreError:
    render_auth_page()
    st.error("Could not start the template database. Please contact the administrator.")
    st.stop()

if not st.session_state.authenticated:
    render_auth_page()
    st.stop()

if "history" not in st.session_state:
    st.session_state.history = []

if "selected_history" not in st.session_state:
    st.session_state.selected_history = None

if "query" not in st.session_state:
    st.session_state.query = ""

if "active_section" not in st.session_state:
    st.session_state.active_section = "workspace"

if "documents_view" not in st.session_state:
    st.session_state.documents_view = "document"

if "active_project_id" not in st.session_state:
    st.session_state.active_project_id = None

prompt_store.initialize_prompt_store()

NAVIGATION_TO_SECTION = {
    "Dashboard": "workspace",
    "Documents": "documents",
    "Agents": "ai",
    "Automations": "automations",
    "Logs": "logs",
    "Settings": "settings",
}
SECTION_TO_NAVIGATION = {
    section: navigation for navigation, section in NAVIGATION_TO_SECTION.items()
}

if "nav_selected" not in st.session_state:
    st.session_state.nav_selected = SECTION_TO_NAVIGATION.get(
        st.session_state.active_section,
        "Dashboard",
    )

if "last_nav_selected" not in st.session_state:
    st.session_state.last_nav_selected = st.session_state.nav_selected

with st.sidebar:
    st.markdown("## Gao Intelligence Hub")

    selected = st.radio(
        "Navigation",
        [
            "Dashboard",
            "Documents",
            "Agents",
            "Automations",
            "Logs",
            "Settings",
        ],
        key="nav_selected",
    )

    if selected != st.session_state.last_nav_selected:
        set_active_section(NAVIGATION_TO_SECTION[selected])
        st.session_state.last_nav_selected = selected
        st.rerun()

    current_user = st.session_state.current_user or {}
    st.caption(f"Logged in as: {st.session_state.username or current_user.get('username', 'User')}")
    st.caption(f"Role: {st.session_state.role or current_user.get('role', 'User')}")
    render_user_limit_summary()

    if st.button("Sign Out", use_container_width=True):
        clear_authenticated_user()
        st.rerun()

main_column, copilot_column = st.columns([3, 1])

with main_column:
    active_section = st.session_state.active_section
    if active_section == "workspace":
        render_workspace()
    elif active_section == "documents":
        render_documents_page()
    elif active_section == "projects":
        render_projects_page()
    elif active_section == "automations":
        render_automation_page()
    elif active_section == "marketplace":
        render_marketplace()
    elif active_section == "ai":
        render_prompt_builder()
    elif active_section == "logs":
        render_logs()
    elif active_section == "settings":
        render_settings()
    else:
        st.session_state.active_section = "workspace"
        render_workspace()

with copilot_column:
    render_ai_copilot_panel()
